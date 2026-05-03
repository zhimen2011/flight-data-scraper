"""
Flight Analysis Web GUI Server
Start with: python flight_server.py → http://localhost:8765
"""
import os, sys, json, csv, glob, threading, webbrowser, io, urllib.parse, re
from datetime import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from analyze_flight import (
    CountryIndex, GEOJSON_PATH, CSV_DIR, analyze_flight, find_csv_files, load_config,
)
from flight_data_scraper import (
    get_flight_plan_points, get_flight_his_pos, save_csv,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORTS_DIR = os.path.join(BASE_DIR, "reports")
SCRAPE_URL = "http://192.168.8.18:8082"
PROMPT_PATH = os.path.join(BASE_DIR, "report_prompt.txt")
CONFIG_PATH = os.path.join(BASE_DIR, "deepseek_config.json")
os.makedirs(REPORTS_DIR, exist_ok=True)

GUI_HTML = r"""
<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="UTF-8"><title>航班轨迹偏差分析工具</title>
<style>
*{margin:0;padding:0;box-sizing:border-box;}
body{font-family:"Microsoft YaHei","SimHei",sans-serif;background:#1a1a2e;color:#e0e0e0;height:100vh;display:flex;flex-direction:column;}
.header{background:#16213e;padding:10px 20px;border-bottom:2px solid #0f3460;}
.header h1{font-size:18px;color:#e94560;}
.tabs{display:flex;background:#0f3460;border-bottom:1px solid #1a1a2e;}
.tab{padding:10px 20px;cursor:pointer;font-size:13px;border-bottom:2px solid transparent;transition:.2s;}
.tab:hover{background:#1a1a3e;}.tab.active{border-bottom-color:#e94560;color:#e94560;}
.tab-content{flex:1;overflow-y:auto;padding:12px 16px;display:none;}
.tab-content.active{display:block;}
.search-bar{display:flex;gap:8px;align-items:center;flex-wrap:wrap;margin-bottom:10px;}
.search-bar input,select{background:#16213e;border:1px solid #333;color:#eee;padding:6px 10px;border-radius:4px;font-size:12px;}
.search-bar button{background:#e94560;color:#fff;border:none;padding:6px 14px;border-radius:4px;cursor:pointer;font-size:12px;}
button:hover{opacity:0.85;}
table{width:100%;border-collapse:collapse;font-size:12px;margin:8px 0;}
th{background:#0f3460;color:#e94560;padding:6px 8px;text-align:left;position:sticky;top:0;}
td{padding:5px 8px;border-bottom:1px solid #1a1a2e;}
tr:hover td{background:#1a1a3e;}
input[type=checkbox]{accent-color:#e94560;transform:scale(1.2);}
.btn-row{display:flex;gap:8px;margin:8px 0;flex-wrap:wrap;align-items:center;}
.btn-row button,.btn-primary,.btn-secondary{padding:6px 14px;border-radius:4px;cursor:pointer;font-size:12px;border:none;}
.btn-primary{background:#e94560;color:#fff;}.btn-secondary{background:#16213e;color:#ccc;border:1px solid #333;}
.progress{width:100%;height:6px;background:#0f3460;border-radius:3px;margin:8px 0;overflow:hidden;}
.progress-bar{height:100%;background:#e94560;width:0%;transition:width .3s;}
.status{font-size:12px;color:#aaa;margin:4px 0;}
.config-row{display:flex;gap:20px;align-items:center;flex-wrap:wrap;margin:8px 0;}
.config-row label{font-size:12px;color:#aaa;}
.config-row input[type=range]{width:140px;accent-color:#e94560;}
.val{color:#e94560;font-weight:bold;}
.section-box{background:#16213e;border:1px solid #0f3460;border-radius:6px;padding:10px 14px;margin:8px 0;}
.section-box h3{font-size:13px;color:#e94560;margin-bottom:6px;cursor:pointer;}
.section-box h3 span{font-size:11px;color:#888;}
.section-box textarea{width:100%;height:200px;background:#0f3460;color:#ccc;border:1px solid #333;border-radius:4px;font-size:11px;font-family:Consolas,monospace;padding:8px;resize:vertical;}
.section-box textarea:focus{outline:none;border-color:#e94560;}
.upload-row{display:flex;gap:8px;align-items:center;margin:6px 0;}
.upload-row input[type=file]{display:none;}
.upload-row .file-label{background:#0f3460;color:#ccc;border:1px solid #333;padding:4px 10px;border-radius:4px;font-size:11px;cursor:pointer;}
.upload-row .file-label:hover{border-color:#e94560;}
.upload-row .file-name{font-size:11px;color:#888;max-width:200px;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;}
select{background:#16213e;border:1px solid #333;color:#eee;padding:3px 6px;border-radius:3px;font-size:11px;}
select:focus{outline:none;border-color:#e94560;}
.dot-ok{color:#4caf50;}.dot-fail{color:#f44336;}.dot-unknown{color:#ffaa00;}
</style></head><body>
<div class="header"><h1>✈ 航班轨迹偏差分析工具</h1></div>
<div class="tabs">
  <div class="tab active" onclick="switchTab(0)">1. 数据抓取</div>
  <div class="tab" onclick="switchTab(1)">2. 偏差分析</div>
  <div class="tab" onclick="switchTab(2)">3. 报告查看</div>
</div>

<!-- TAB 0: Scrape -->
<div class="tab-content active" id="tab0">
  <div class="search-bar">
    <span>航班号:</span><input id="sFlight" placeholder="I98831,I98833" style="width:130px;">
    <span>飞机号:</span><input id="sAircraft" placeholder="B-2079" style="width:80px;">
    <span>日期:</span><input id="sDateFrom" value="2026-04-01" style="width:95px;">
    <span>~</span><input id="sDateTo" value="2026-04-30" style="width:95px;">
    <button onclick="searchFlights()">搜索并添加到结果池</button>
  </div>
  <div class="status" id="baseInfo" style="display:none;"></div>
  <div class="btn-row">
    <button class="btn-secondary" onclick="selPool('outbound')">全选去程</button>
    <button class="btn-secondary" onclick="selPool('return')">全选回程</button>
    <button class="btn-secondary" onclick="selPool('unknown')">全选未知</button>
    <button class="btn-secondary" onclick="selPool('all')">全选</button>
    <button class="btn-secondary" onclick="selPool('none')">清空</button>
    <button class="btn-secondary" onclick="removeUnchecked()">移除未勾选</button>
  </div>
  <div class="status" id="poolStatus">结果池（共 0 条）</div>
  <div id="poolTable"></div>
  <div class="btn-row">
    <button class="btn-primary" onclick="startScrape()">开始抓取勾选的航班</button>
    <div class="progress" style="flex:1;min-width:200px;"><div class="progress-bar" id="scrapeProgress"></div></div>
    <span class="status" id="scrapeStatus">就绪</span>
  </div>
</div>

<!-- TAB 1: Analyze -->
<div class="tab-content" id="tab1">
  <!-- API Key -->
  <div class="section-box">
    <h3 onclick="toggleSection('apiSection')">🔑 API 配置 <span id="apiStatus" class="dot-unknown">● 未检测</span></h3>
    <div id="apiSection">
      <div class="config-row">
        <span>Key:</span><input id="apiKey" style="width:300px;" placeholder="sk-...">
        <span>Model:</span><input id="apiModel" value="deepseek-v4-pro" style="width:140px;">
        <button class="btn-secondary" onclick="testApiKey()">测试连接</button>
        <button class="btn-secondary" onclick="saveApiKey()">保存</button>
      </div>
    </div>
  </div>

  <!-- Thresholds -->
  <div class="config-row">
    <label>高度偏差阈值: <input type="range" id="aAlt" min="200" max="3000" value="1000" step="100" oninput="$('aAltVal').textContent=$('aAlt').value">
      <span class="val" id="aAltVal">1000</span> ft</label>
    <label>持续时间阈值: <input type="range" id="aDur" min="20" max="500" value="200" step="10" oninput="$('aDurVal').textContent=$('aDur').value">
      <span class="val" id="aDurVal">200</span> nm</label>
    <label><input type="checkbox" id="aCruise" checked> 仅巡航段</label>
    <button class="btn-secondary" onclick="saveConfig()">保存为默认</button>
  </div>

  <!-- CSV list -->
  <div class="btn-row">
    <button class="btn-secondary" onclick="selAnalyze(true)">全选</button>
    <button class="btn-secondary" onclick="selAnalyze(false)">清空</button>
    <button class="btn-secondary" onclick="refreshAnalyze()">刷新列表</button>
  </div>
  <div class="status" id="analyzeStatus">已有数据</div>
  <div id="analyzeTable"></div>

  <!-- Action buttons -->
  <div class="btn-row">
    <button class="btn-primary" onclick="runAnalysis()">▶ 运行偏差分析</button>
    <button class="btn-primary" onclick="genWordReport()">生成 Word 报告</button>
    <button class="btn-primary" onclick="genHtmlReport()">生成 HTML 集成报告</button>
    <button class="btn-secondary" onclick="runAll()">全部（分析+Word+HTML）</button>
    <div class="progress" style="flex:1;min-width:150px;"><div class="progress-bar" id="analyzeProgress"></div></div>
    <span class="status" id="analyzeMsg">就绪</span>
  </div>

  <!-- Prompt editor -->
  <div class="section-box">
    <h3 onclick="toggleSection('promptSection')">📝 报告 Prompt 模板 <span>(点击展开/折叠)</span></h3>
    <div id="promptSection" style="display:none;">
      <div class="upload-row">
        <span>参考主报告:</span><input type="file" id="refMain" accept=".docx" onchange="uploadRef('main')">
        <label class="file-label" for="refMain">选择文件</label>
        <span class="file-name" id="refMainName"></span>
      </div>
      <div class="upload-row">
        <span>参考附件:</span><input type="file" id="refApp" accept=".docx" onchange="uploadRef('app')">
        <label class="file-label" for="refApp">选择文件</label>
        <span class="file-name" id="refAppName"></span>
      </div>
      <textarea id="promptText" placeholder="加载中..."></textarea>
      <div class="btn-row">
        <button class="btn-secondary" onclick="loadPrompt()">刷新</button>
        <button class="btn-secondary" onclick="resetPrompt()">恢复默认</button>
        <button class="btn-primary" onclick="savePrompt()">保存 Prompt</button>
        <span class="status" id="promptMsg"></span>
      </div>
    </div>
  </div>
</div>

<!-- TAB 2: Reports -->
<div class="tab-content" id="tab2">
  <div class="btn-row">
    <button class="btn-secondary" onclick="refreshReports()">刷新</button>
    <button class="btn-secondary" onclick="openReportsFolder()">打开文件夹</button>
  </div>
  <div id="reportsTable"></div>
</div>

<script>
function switchTab(n){
  document.querySelectorAll('.tab').forEach((t,i)=>t.classList.toggle('active',i===n));
  document.querySelectorAll('.tab-content').forEach((t,i)=>t.classList.toggle('active',i===n));
  if(n===1){refreshAnalyze();loadApiCfg();}
  if(n===2) refreshReports();
}
function toggleSection(id){var e=$(id);e.style.display=e.style.display==='none'?'block':'none';}
async function api(path,opt={}){
  try{
    let url='/api'+path,r;
    if(opt.method==='POST'){r=await fetch(url,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(opt.body||{})});}
    else{r=await fetch(url);}
    if(!r.ok){let e=await r.text();throw new Error(e||r.statusText);}
    return await r.json();
  }catch(e){alert('API 错误: '+e.message);return null;}
}
async function apiForm(path,formData){
  try{let r=await fetch('/api'+path,{method:'POST',body:formData});if(!r.ok)throw new Error(await r.text());return await r.json();}
  catch(e){alert('API 错误: '+e.message);return null;}
}
function $(id){return document.getElementById(id);}
function setStatus(id,txt){$(id).textContent=txt;}
function setProgress(id,pct){$(id).style.width=pct+'%';}
async function readFileAsText(file){return new Promise((resolve)=>{let r=new FileReader();r.onload=()=>resolve(r.result);r.readAsText(file);});}

// ── TAB 0: Pool ──
let pool=[],poolChecks={},autoBase='';
async function searchFlights(){
  let qs=`?flight=${encodeURIComponent($('sFlight').value)}&aircraft=${encodeURIComponent($('sAircraft').value)}&from=${$('sDateFrom').value}&to=${$('sDateTo').value}`;
  let r=await api('/search'+qs);
  if(!r){return;} if(r.error){alert(r.error);return;}
  let added=0;
  r.results.forEach(fl=>{let k=fl.key;if(!pool.some(p=>p.key===k)){pool.push(fl);poolChecks[k]=true;added++;}});
  if(r.autoBase){autoBase=r.autoBase;$('baseInfo').style.display='block';
    $('baseInfo').textContent='自动识别基地: '+autoBase+' (从此机场起飞=去程，降落=回程)';}
  renderPool();setStatus('poolStatus','结果池（共 '+pool.length+' 条）');alert('新增 '+added+' 条');
}
function renderPool(){
  let h='<table><tr><th>✓</th><th>方向</th><th>日期</th><th>航班号</th><th>机号</th><th>起降</th></tr>';
  pool.forEach(p=>{
    let ck=poolChecks[p.key]!==false?'☑':'☐';
    let selId='dir_'+p.key.replace(/[^a-zA-Z0-9]/g,'_');
    h+=`<tr><td onclick="togglePool('${p.key}')" style="cursor:pointer">${ck}</td>
      <td><select id="${selId}" onchange="changeDir('${p.key}',this.value)" style="width:65px;">
        <option value="去程"${p.dir==='去程'?' selected':''}>去程</option>
        <option value="回程"${p.dir==='回程'?' selected':''}>回程</option>
        <option value="未知"${p.dir==='未知'?' selected':''}>未知</option>
      </select></td>
      <td>${p.date}</td><td>${p.flight}</td><td>${p.ac}</td><td>${p.route}</td></tr>`;
  });
  h+='</table>';$('poolTable').innerHTML=h;
}
function changeDir(key,val){let p=pool.find(x=>x.key===key);if(p)p.dir=val;}
function togglePool(k){poolChecks[k]=!poolChecks[k];renderPool();}
function selPool(mode){
  pool.forEach(p=>{
    if(mode==='all') poolChecks[p.key]=true;
    else if(mode==='none') poolChecks[p.key]=false;
    else if(mode==='outbound') poolChecks[p.key]=p.dir==='去程';
    else if(mode==='return') poolChecks[p.key]=p.dir==='回程';
    else if(mode==='unknown') poolChecks[p.key]=p.dir==='未知';
  });
  renderPool();
}
function removeUnchecked(){pool=pool.filter(p=>poolChecks[p.key]!==false);renderPool();setStatus('poolStatus','结果池（共 '+pool.length+' 条）');}
async function startScrape(){
  let flights=pool.filter(p=>poolChecks[p.key]!==false);
  if(!flights.length){alert('请勾选航班');return;}
  setStatus('scrapeStatus','抓取中...');setProgress('scrapeProgress',0);
  let r=await api('/scrape',{method:'POST',body:{flights}});
  setProgress('scrapeProgress',100);if(!r){setStatus('scrapeStatus','失败');return;}
  setStatus('scrapeStatus',r.status);
  if(r.results){alert(r.status+'\n'+r.results.map(x=>x.flight+' '+x.date+': '+(x.ok?'OK '+x.detail:'FAIL '+x.detail)).join('\n'));}
  setTimeout(refreshAnalyze,500);
}

// ── TAB 1: API Config ──
async function loadApiCfg(){
  let r=await api('/load_api_cfg');
  if(!r)return;$('apiKey').value=r.api_key||'';$('apiModel').value=r.model||'deepseek-v4-pro';
  if(r.status==='ok'){$('apiStatus').innerHTML='<span class="dot-ok">● 正常</span>';}
  else if(r.status==='unset'){$('apiStatus').innerHTML='<span class="dot-unknown">● 未配置</span>';}
  else{$('apiStatus').innerHTML='<span class="dot-fail">● '+r.status+'</span>';}
}
async function testApiKey(){
  $('apiStatus').innerHTML='<span class="dot-unknown">● 检测中...</span>';
  let r=await api('/test_key',{method:'POST',body:{key:$('apiKey').value,model:$('apiModel').value}});
  if(!r){$('apiStatus').innerHTML='<span class="dot-fail">● 网络错误</span>';return;}
  if(r.ok){$('apiStatus').innerHTML='<span class="dot-ok">● 正常</span>';alert('API Key 有效');}
  else{$('apiStatus').innerHTML='<span class="dot-fail">● 失效</span>';alert('API Key 无效: '+r.error);}
}
async function saveApiKey(){
  let r=await api('/save_key',{method:'POST',body:{key:$('apiKey').value,model:$('apiModel').value}});
  if(r)alert('已保存');loadApiCfg();
}

// ── TAB 1: Analyze ──
let analyzeList=[],analyzeChecks={};
async function refreshAnalyze(){
  let r=await api('/list_csv');if(!r)return;
  analyzeList=r.files||[];analyzeList.forEach(f=>{if(!(f.key in analyzeChecks))analyzeChecks[f.key]=true;});
  renderAnalyze();setStatus('analyzeStatus','共 '+analyzeList.length+' 个航班数据');
}
function renderAnalyze(){
  let h='<table><tr><th>✓</th><th>日期</th><th>航班号</th><th>状态</th></tr>';
  analyzeList.forEach(f=>{
    let ck=analyzeChecks[f.key]!==false?'☑':'☐';
    h+=`<tr><td onclick="toggleAnalyze('${f.key}')" style="cursor:pointer">${ck}</td><td>${f.date}</td><td>${f.flight}</td><td>${f.status||'CSV ✓'}</td></tr>`;
  });
  h+='</table>';$('analyzeTable').innerHTML=h;
}
function toggleAnalyze(k){analyzeChecks[k]=!analyzeChecks[k];renderAnalyze();}
function selAnalyze(state){analyzeList.forEach(f=>analyzeChecks[f.key]=state);renderAnalyze();}
function updateCfg(){$('aAltVal').textContent=$('aAlt').value;$('aDurVal').textContent=$('aDur').value;}
async function saveConfig(){
  let cfg={min_alt_deviation_ft:parseInt($('aAlt').value),min_duration_nm:parseInt($('aDur').value)};
  await api('/save_config',{method:'POST',body:cfg});alert('配置已保存');
}
async function runAnalysis(){
  let keys=analyzeList.filter(f=>analyzeChecks[f.key]!==false).map(f=>f.key);
  if(!keys.length){alert('请勾选航班');return;}
  let cfg={min_alt_deviation_ft:parseInt($('aAlt').value),min_duration_nm:parseInt($('aDur').value)};
  setStatus('analyzeMsg','分析中...');setProgress('analyzeProgress',0);
  let r=await api('/analyze',{method:'POST',body:{keys,config:cfg}});
  setProgress('analyzeProgress',100);if(!r){setStatus('analyzeMsg','分析失败');return;}
  setStatus('analyzeMsg',r.status||('完成 '+r.count+' 个航班'));
}
async function genWordReport(){
  setStatus('analyzeMsg','生成 Word 报告...');setProgress('analyzeProgress',0);
  let r=await api('/gen_word',{method:'POST',body:{}});
  setProgress('analyzeProgress',100);if(!r){setStatus('analyzeMsg','Word 报告失败');return;}
  setStatus('analyzeMsg',r.status||'Word 报告已生成');refreshReports();
}
async function genHtmlReport(){
  setStatus('analyzeMsg','生成 HTML 报告...');setProgress('analyzeProgress',0);
  let r=await api('/gen_html',{method:'POST',body:{}});
  setProgress('analyzeProgress',100);if(!r){setStatus('analyzeMsg','HTML 报告失败');return;}
  setStatus('analyzeMsg',r.status||'HTML 报告已生成');refreshReports();
}
async function runAll(){await runAnalysis();setTimeout(genWordReport,2000);setTimeout(genHtmlReport,4000);}

// ── TAB 1: Prompt Editor ──
async function loadPrompt(){let r=await api('/load_prompt');if(r){$('promptText').value=r.prompt||'';setStatus('promptMsg','已加载');}}
async function savePrompt(){
  let r=await api('/save_prompt',{method:'POST',body:{prompt:$('promptText').value}});
  if(r)setStatus('promptMsg','已保存');
}
async function resetPrompt(){let r=await api('/reset_prompt');if(r){$('promptText').value=r.prompt||'';setStatus('promptMsg','已恢复默认');}}
async function uploadRef(type){
  let inp=$(type==='main'?'refMain':'refApp');
  if(!inp.files[0])return;
  let name=inp.files[0].name;
  $(type==='main'?'refMainName':'refAppName').textContent=name;
  let fd=new FormData();fd.append('file',inp.files[0]);fd.append('type',type);
  let r=await apiForm('/upload_ref',fd);
  if(r){setStatus('promptMsg','参考文件已上传');loadPrompt();}
}

// ── TAB 2: Reports ──
async function refreshReports(){
  let r=await api('/list_reports');if(!r)return;
  let h='<table><tr><th>文件名</th><th>大小</th><th>时间</th><th>操作</th></tr>';
  (r.files||[]).forEach(f=>{h+=`<tr><td>${f.name}</td><td>${f.size}</td><td>${f.time}</td><td><a href="/reports/${encodeURIComponent(f.name)}" target="_blank" style="color:#e94560">打开</a></td></tr>`;});
  h+='</table>';$('reportsTable').innerHTML=h;
}
function openReportsFolder(){window.open('/reports/');}

// Init
loadApiCfg();loadPrompt();refreshAnalyze();refreshReports();
</script></body></html>
""";

# ─── HTTP Server ──────────────────────────────────────────────────────────────

class FlightAPIHandler(BaseHTTPRequestHandler):
    search_pool = []
    analysis_results = {}
    country_index = None
    ref_main_text = ""
    ref_app_text = ""

    def log_message(self, format, *args): pass

    def do_GET(self):
        path = self.path.split("?")[0]
        if path in ("/", "/index.html"): self._serve_html(GUI_HTML)
        elif path == "/api/search": self._api_search()
        elif path == "/api/list_csv": self._api_list_csv()
        elif path == "/api/list_reports": self._api_list_reports()
        elif path == "/api/load_api_cfg": self._api_load_api_cfg()
        elif path == "/api/load_prompt": self._api_load_prompt()
        elif path == "/api/reset_prompt": self._api_reset_prompt()
        elif path.startswith("/reports/"): self._serve_report_file(path)
        else: self._send_json({"error": "not found"}, 404)

    def do_POST(self):
        body = self._read_body()
        path = self.path.split("?")[0]

        # multipart upload
        ctype = self.headers.get("Content-Type", "")
        if "multipart/form-data" in ctype and path == "/api/upload_ref":
            self._api_upload_ref()
            return

        data = json.loads(body) if body else {}
        if path == "/api/scrape": self._api_scrape(data)
        elif path == "/api/analyze": self._api_analyze(data)
        elif path == "/api/gen_word": self._api_gen_word()
        elif path == "/api/gen_html": self._api_gen_html()
        elif path == "/api/save_config": self._save_config(data); self._send_json({"status":"ok"})
        elif path == "/api/test_key": self._api_test_key(data)
        elif path == "/api/save_key": self._api_save_key(data)
        elif path == "/api/save_prompt": self._api_save_prompt(data)
        else: self._send_json({"error": "not found"}, 404)

    def _read_body(self):
        length = int(self.headers.get("Content-Length", 0))
        return self.rfile.read(length).decode("utf-8") if length else ""

    def _send_json(self, data, code=200):
        self.send_response(code)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.end_headers()
        self.wfile.write(json.dumps(data, ensure_ascii=False).encode("utf-8"))

    def _serve_html(self, html):
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        for h in ("Cache-Control","Pragma","Expires"):
            self.send_header(h, "no-cache")
        self.end_headers()
        self.wfile.write(html.encode("utf-8"))

    def _serve_report_file(self, path):
        fname = os.path.basename(urllib.parse.unquote(path))
        fpath = os.path.join(REPORTS_DIR, fname)
        if os.path.isfile(fpath):
            self.send_response(200)
            ct = "text/html" if fname.endswith(".html") else "application/octet-stream"
            self.send_header("Content-Type", ct)
            self.end_headers()
            with open(fpath, "rb") as f: self.wfile.write(f.read())
        else: self._send_json({"error":"not found"}, 404)

    # ── API: Config ──
    def _api_load_api_cfg(self):
        if os.path.exists(CONFIG_PATH):
            with open(CONFIG_PATH) as f: cfg = json.load(f)
            cfg["status"] = "ok"
            self._send_json(cfg)
        else: self._send_json({"status":"unset","api_key":"","model":"deepseek-v4-pro"})

    def _api_test_key(self, data):
        import requests as req
        key = data.get("key",""); model = data.get("model","deepseek-v4-pro")
        if not key: self._send_json({"ok":False,"error":"Key is empty"}); return
        try:
            r = req.post("https://api.deepseek.com/v1/chat/completions",
                headers={"Authorization":f"Bearer {key}","Content-Type":"application/json"},
                json={"model":model,"messages":[{"role":"user","content":"hi"}],"max_tokens":5}, timeout=15)
            if r.status_code == 200: self._send_json({"ok":True})
            else: self._send_json({"ok":False,"error":f"HTTP {r.status_code}: {r.text[:100]}"})
        except Exception as e: self._send_json({"ok":False,"error":str(e)[:100]})

    def _api_save_key(self, data):
        cfg = {}
        if os.path.exists(CONFIG_PATH):
            with open(CONFIG_PATH) as f: cfg = json.load(f)
        cfg["api_key"] = data.get("key","")
        cfg["model"] = data.get("model","deepseek-v4-pro")
        with open(CONFIG_PATH,"w") as f: json.dump(cfg, f, indent=2)
        self._send_json({"status":"ok"})

    # ── API: Search ──
    def _api_search(self):
        qs = urllib.parse.urlparse(self.path).query
        params = urllib.parse.parse_qs(qs)
        flight_nums = [f.strip().upper() for f in params.get("flight",[""])[0].split(",") if f.strip()]
        aircraft = params.get("aircraft",[""])[0].strip().upper()
        date_from = params.get("from",[""])[0]; date_to = params.get("to",[""])[0]
        if not flight_nums and not aircraft: self._send_json({"error":"请输入航班号或飞机号"}); return

        import requests as req
        results = []; dep_counts = {}
        search_terms = flight_nums if flight_nums else [aircraft]
        for fn in search_terms:
            try:
                r = req.get(f"{SCRAPE_URL}/getFlightListByAN",
                    params={"fi":fn,"staDate":date_from,"endDate":date_to}, timeout=30)
                data = r.json()
                if isinstance(data, dict): data = data.get("data",[])
                for fl in data:
                    fid = fl.get("fLIGHTID",""); tko = fl.get("tKO_TIME","")[:10]
                    ac = fl.get("aIRCRAFT","")
                    if aircraft and aircraft not in ac.upper(): continue
                    dep = fl.get("tKO_FIELD",""); arr = fl.get("dES_FIELD","")
                    dep_counts[dep] = dep_counts.get(dep,0) + 1
                    key = f"{tko}_{fid}"
                    results.append({"key":key,"dir":"?","date":tko,"flight":fn,"ac":ac,
                        "route":f"{dep}→{arr}","flightId":fid,"aircraft":ac,
                        "tkoTime":fl.get("tKO_TIME",""),"desTime":fl.get("dES_TIME",""),
                        "tkoTimeOff":fl.get("tKO_TIME_OFF",""),"depAirport":dep,"arrAirport":arr})
            except Exception as e: print(f"  search error: {e}")

        # Auto-detect base airport
        auto_base = ""
        if dep_counts:
            auto_base = max(dep_counts, key=dep_counts.get)
            for r2 in results:
                if r2["dir"] == "?":
                    if r2["depAirport"] == auto_base: r2["dir"] = "去程"
                    elif r2["arrAirport"] == auto_base: r2["dir"] = "回程"
                    else: r2["dir"] = "未知"
        self._send_json({"results":results,"autoBase":auto_base})

    # ── API: Scrape ──
    def _api_scrape(self, data):
        flights = data.get("flights",[])
        if not flights: self._send_json({"status":"无航班"}); return
        import requests as req
        csv_dir = os.path.join(BASE_DIR,"flight_csv"); os.makedirs(csv_dir,exist_ok=True)
        results = []
        for fl_info in flights:
            s = {"flight":fl_info.get("flight","?"),"date":fl_info.get("date","?"),"ok":False,"detail":""}
            try:
                fid=fl_info.get("flightId","");fn=fl_info.get("flight","");ac=fl_info.get("aircraft","")
                dep=fl_info.get("depAirport","");arr=fl_info.get("arrAirport","")
                tko=fl_info.get("tkoTime","");des=fl_info.get("desTime","")
                tko_off=fl_info.get("tkoTimeOff","");fdate=tko[:10] if tko else fl_info.get("date","")
                prefix=os.path.join(csv_dir,f"{fn}_{fdate}");details=[]
                plan=get_flight_plan_points(fid,ac,dep,arr,tko_off or tko)
                if plan and len(plan)>0:
                    save_csv(plan,f"{prefix}_plan_track.csv")
                    save_csv(plan,f"{prefix}_plan_profile.csv",["name","dist","alt","ful","time","lat","lon"])
                    details.append(f"plan={len(plan)}pts")
                else: details.append("plan=empty")
                actual=get_flight_his_pos(fid,tko,des)
                ad=actual
                if isinstance(actual,dict): ad=actual.get("data",[])
                if ad and len(ad)>0:
                    save_csv(ad,f"{prefix}_actual_track.csv")
                    save_csv(ad,f"{prefix}_actual_profile.csv",["gateway_time","alt","fob","dis","lat","lon","posPointName","sAlt","type"])
                    details.append(f"actual={len(ad)}pts")
                else: details.append("actual=empty")
                s["ok"]=True;s["detail"]=", ".join(details)
            except Exception as e: s["detail"]=str(e)[:100]
            results.append(s)
        ok_count=sum(1 for r in results if r["ok"])
        self._send_json({"status":f"完成 {ok_count}/{len(results)} 个航班","results":results})

    # ── API: List ──
    def _api_list_csv(self):
        files=[];csv_dir=os.path.join(BASE_DIR,"flight_csv")
        for root,dirs,fnames in os.walk(csv_dir):
            for f in fnames:
                if f.endswith("_plan_track.csv"):
                    base=f.replace("_plan_track.csv","");parts=base.split("_")
                    if len(parts)>=2: files.append({"key":f"{parts[0]}_{parts[1]}","flight":parts[0],"date":parts[1],"status":"CSV ✓"})
        files.sort(key=lambda x:(x["date"],x["flight"]),reverse=True)
        self._send_json({"files":files})

    def _api_list_reports(self):
        files=[]
        if os.path.exists(REPORTS_DIR):
            for f in sorted(os.listdir(REPORTS_DIR),reverse=True):
                p=os.path.join(REPORTS_DIR,f)
                if os.path.isfile(p) and not f.startswith("~$"):
                    files.append({"name":f,"size":f"{os.path.getsize(p)/1024:.0f} KB","time":datetime.fromtimestamp(os.path.getmtime(p)).strftime("%Y-%m-%d %H:%M")})
        self._send_json({"files":files})

    # ── API: Analyze ──
    def _api_analyze(self, data):
        keys=data.get("keys",[]);config=data.get("config",{"min_alt_deviation_ft":1000,"min_duration_nm":200})
        if FlightAPIHandler.country_index is None: FlightAPIHandler.country_index = CountryIndex(GEOJSON_PATH)
        results={}
        for key in keys:
            parts=key.split("_",1);fn,ds=parts[0],parts[1] if len(parts)>1 else ""
            pf,af=find_csv_files(fn,ds)
            if pf and af:
                r=analyze_flight(pf,af,config,FlightAPIHandler.country_index)
                if r: results[key]=r
        FlightAPIHandler.analysis_results=results
        self._send_json({"status":f"完成 {len(results)} 个航班","count":len(results)})

    # ── API: Generate Reports ──
    def _get_sys_prompt(self):
        """Build system prompt with optional reference docs."""
        from review_report import build_system_prompt
        return build_system_prompt(FlightAPIHandler.ref_main_text, FlightAPIHandler.ref_app_text)

    def _api_gen_word(self):
        if not FlightAPIHandler.analysis_results: self._send_json({"status":"请先运行分析"}); return
        from review_report import finalize_to_docx, build_user_prompt, call_deepseek, load_api_config
        analysis={"flights":FlightAPIHandler.analysis_results,"analysis_time":datetime.now().isoformat()}
        api_cfg=load_api_config();rn="航线分析"
        for k in FlightAPIHandler.analysis_results: rn=k.split("_")[0];break
        dl=sorted(set(k.rsplit("_",1)[1] for k in FlightAPIHandler.analysis_results))
        sp=self._get_sys_prompt()
        up=build_user_prompt(analysis,{"route_name":rn,"date_range":f"{dl[0]}至{dl[-1]}"})
        try:
            api_resp=call_deepseek([{"role":"system","content":sp},{"role":"user","content":up}],api_cfg)
        except Exception as e: self._send_json({"status":f"API 错误: {e}"}); return
        ts=datetime.now().strftime("%Y%m%d_%H%M")
        dp=os.path.join(REPORTS_DIR,f"{rn}_draft_{ts}.md")
        with open(dp,"w",encoding="utf-8") as f: f.write(api_resp)
        finalize_to_docx(dp,analysis,dp.replace(".md",".docx"))
        self._send_json({"status":"Word 报告已生成"})

    def _api_gen_html(self):
        if not FlightAPIHandler.analysis_results: self._send_json({"status":"请先运行分析"}); return
        from review_report import generate_integrated_html, build_user_prompt, call_deepseek, load_api_config
        analysis={"flights":FlightAPIHandler.analysis_results,"analysis_time":datetime.now().isoformat()}
        api_cfg=load_api_config();rn="航班分析"
        for k in FlightAPIHandler.analysis_results: rn=k.split("_")[0];break
        dl=sorted(set(k.rsplit("_",1)[1] for k in FlightAPIHandler.analysis_results))
        config={"min_alt_deviation_ft":1000,"min_duration_nm":200}
        sp=self._get_sys_prompt()
        up=build_user_prompt(analysis,{"route_name":rn,"date_range":f"{dl[0]}至{dl[-1]}"})
        try:
            api_resp=call_deepseek([{"role":"system","content":sp},{"role":"user","content":up}],api_cfg)
        except Exception as e: self._send_json({"status":f"API 错误: {e}"}); return
        ts=datetime.now().strftime("%Y%m%d_%H%M")
        hp=os.path.join(REPORTS_DIR,f"{rn}_integrated_{ts}.html")
        generate_integrated_html(analysis,api_resp,config,hp)
        self._send_json({"status":"HTML 报告已生成"})

    # ── API: Prompt ──
    def _api_load_prompt(self):
        prompt = ""
        if os.path.exists(PROMPT_PATH):
            with open(PROMPT_PATH,"r",encoding="utf-8") as f: prompt = f.read()
        if not prompt:
            from review_report import build_system_prompt
            prompt = build_system_prompt()
        self._send_json({"prompt":prompt})

    def _api_save_prompt(self, data):
        prompt = data.get("prompt","")
        with open(PROMPT_PATH,"w",encoding="utf-8") as f: f.write(prompt)
        self._send_json({"status":"ok"})

    def _api_reset_prompt(self):
        if os.path.exists(PROMPT_PATH): os.remove(PROMPT_PATH)
        from review_report import build_system_prompt
        self._send_json({"prompt":build_system_prompt()})

    def _api_upload_ref(self):
        """Handle multipart file upload for reference docs."""
        ctype = self.headers.get("Content-Type","")
        length = int(self.headers.get("Content-Length",0))
        raw = self.rfile.read(length)

        boundary = None
        for part in ctype.split(";"):
            if "boundary=" in part:
                boundary = part.split("boundary=")[1].strip().strip('"')
                break
        if not boundary: self._send_json({"status":"no boundary"}); return

        # Parse type field
        ref_type = "main"
        ref_content = None
        for segment in raw.split(b"--"+boundary.encode()):
            if b"Content-Disposition" not in segment: continue
            header_end = segment.find(b"\r\n\r\n")
            if header_end < 0: continue
            headers = segment[:header_end].decode("utf-8","ignore")
            content = segment[header_end+4:]
            content = content.rstrip(b"\r\n--\r\n").rstrip(b"\r\n--").rstrip(b"\r\n")

            if 'name="type"' in headers:
                ref_type = content.decode("utf-8","ignore").strip()
            elif 'name="file"' in headers and b"filename=" in segment[:header_end]:
                ref_content = content

        if not ref_content:
            self._send_json({"status":"no file content"}); return

        try:
            from docx import Document
            tmp = os.path.join(BASE_DIR,"_tmp_ref.docx")
            with open(tmp,"wb") as f: f.write(ref_content)
            doc = Document(tmp)
            text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    t = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if t: text += "\n" + t
            os.remove(tmp)
            if ref_type == "app": FlightAPIHandler.ref_app_text = text
            else: FlightAPIHandler.ref_main_text = text
        except Exception as e:
            print(f"  upload_ref error: {e}")
            self._send_json({"status":f"parse error: {e}"}); return

        self._send_json({"status":"ok"})

    def _save_config(self, cfg):
        with open(os.path.join(BASE_DIR,"analysis_config.json"),"w") as f:
            json.dump(cfg,f,indent=2)

# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    os.makedirs(REPORTS_DIR,exist_ok=True)
    port=8765
    server=HTTPServer(("localhost",port),FlightAPIHandler)
    url=f"http://localhost:{port}"
    print(f"\n  航班分析服务器已启动\n  浏览器打开: {url}\n  按 Ctrl+C 停止\n")
    webbrowser.open(url)
    try: server.serve_forever()
    except KeyboardInterrupt: print("\n  服务器已停止"); server.shutdown()

if __name__=="__main__": main()
