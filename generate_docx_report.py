"""
Word Report Generator for Flight Deviation Analysis
====================================================
Generates a .docx main report + appendix with profile charts,
modeled after the manual human-written report format.

Usage:
    python generate_docx_report.py --flight I99806 --dates 2026-04-02,...,2026-04-29
    python generate_docx_report.py --config report_config.json
"""
import os
import sys
import json
import argparse
from datetime import datetime
from collections import defaultdict

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np

# Configure Chinese font for matplotlib
for font_name in ['Microsoft YaHei', 'SimHei', 'SimSun', 'Arial Unicode MS']:
    try:
        fm.findfont(fm.FontProperties(family=font_name), fallback_to_default=False)
        plt.rcParams['font.family'] = font_name
        break
    except Exception:
        continue
plt.rcParams['axes.unicode_minus'] = False  # Fix negative sign display

# Add parent for analyze_flight imports
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from analyze_flight import (
    load_plan_track, load_actual_track, find_csv_files,
    analyze_flight, CountryIndex, GEOJSON_PATH, CSV_DIR,
    DEFAULT_MIN_ALT_DEV_FT, DEFAULT_MIN_DURATION_NM,
)

REPORTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
os.makedirs(REPORTS_DIR, exist_ok=True)

# ─── Flight Level Standards ──────────────────────────────────────────────────
#
# ICAO (international, feet-based):
#   RVSM FL290-FL410, 1000ft separation
#   Eastbound (000°-179°): FL290, 310, 330, 350, 370, 390, 410
#   Westbound (180°-359°): FL300, 320, 340, 360, 380, 400
#   Below FL290: Eastbound=odd thousands (5000,7000,9000,11000...),
#                Westbound=even thousands (6000,8000,10000,12000...)
#   Above FL410: 2000ft separation, odd levels only
#
# CAAC / China (meters-based, converted to FL):
#   RVSM 8900m-12500m, 300m separation
#   Eastbound: 8900m(FL291), 9500m(FL311), 10100m(FL331), 10700m(FL351),
#              11300m(FL371), 11900m(FL391), 12500m(FL411)
#   Westbound: 9200m(FL301), 9800m(FL321), 10400m(FL341), 11000m(FL361),
#              11600m(FL381), 12200m(FL401)
#   Below RVSM: 600m interval in some bands, 300m in others
#
# Key difference: CAAC FLs end in -01/-11/-21... (offset by 100ft from ICAO)

# ── All valid IFR cruising flight levels ──

def _build_all_valid_fls():
    """Build complete set of valid IFR flight levels worldwide."""
    fls = set()
    # ICAO: below FL180 (IFR, 1000ft separation)
    # Eastbound (odd thousands)
    for alt in range(3000, 18000, 2000):
        fls.add(alt // 100)  # FL30, FL50, FL70, FL90, FL110, FL130, FL150, FL170
    # Westbound (even thousands)
    for alt in range(4000, 19000, 2000):
        fls.add(alt // 100)  # FL40, FL60, FL80, FL100, FL120, FL140, FL160, FL180
    # ICAO: FL180-FL290 (above transition, 1000ft)
    for alt in range(18000, 30000, 2000):
        fls.add(alt // 100)  # FL180, FL200, FL220, FL240, FL260, FL280
        fls.add((alt + 1000) // 100)  # FL190, FL210, FL230, FL250, FL270, FL290
    # ICAO RVSM: FL290-FL410
    for fl in [290, 300, 310, 320, 330, 340, 350, 360, 370, 380, 390, 400, 410]:
        fls.add(fl)
    # CAAC RVSM (offset by 1 from ICAO)
    for fl in [291, 301, 311, 321, 331, 341, 351, 361, 371, 381, 391, 401, 411]:
        fls.add(fl)
    # Above FL410: 2000ft separation
    for fl in range(430, 600, 20):
        fls.add(fl)
    return sorted(fls)


ALL_VALID_FL = _build_all_valid_fls()

# CAAC metric RVSM levels: {meters: FL_equivalent}
# Eastbound: 8900(FL291), 9500(FL311), 10100(FL331), 10700(FL351),
#            11300(FL371), 11900(FL391), 12500(FL411)
# Westbound: 9200(FL301), 9800(FL321), 10400(FL341), 11000(FL361),
#            11600(FL381), 12200(FL401)
CAAC_RVSM_M_TO_FL = {
    8900: 291, 9200: 301,
    9500: 311, 9800: 321,
    10100: 331, 10400: 341,
    10700: 351, 11000: 361,
    11300: 371, 11600: 381,
    11900: 391, 12200: 401,
    12500: 411,
}


def nearest_valid_fl(alt_ft):
    """Find the nearest valid IFR cruising flight level for the given altitude."""
    fl = int(round(alt_ft / 100))
    if fl in ALL_VALID_FL:
        return fl
    return min(ALL_VALID_FL, key=lambda x: abs(x - fl))


def format_altitude(alt_ft, region=None):
    """Format altitude respecting regional standards (input in FEET)."""
    if alt_ft <= 0:
        return "0"

    fl = nearest_valid_fl(alt_ft)
    alt_m = round(alt_ft * 0.3048)

    # China: show CAAC standard meters + FL
    if region in ("国内段", "CN"):
        nearest_m = min(CAAC_RVSM_M_TO_FL.keys(), key=lambda m: abs(m - alt_m))
        if nearest_m and abs(nearest_m - alt_m) < 200:
            fl_caac = CAAC_RVSM_M_TO_FL[nearest_m]
            return f"{nearest_m}米(FL{fl_caac})"
        return f"{alt_m}米(FL{fl})"

    # International: standard FL format
    return f"FL{fl}"

def format_fl_deviation(plan_alt, actual_alt, region=None):
    """
    Describe altitude deviation using valid flight levels.
    """
    plan_fl = nearest_valid_fl(plan_alt)
    actual_fl = nearest_valid_fl(actual_alt)
    diff_levels = abs(plan_fl - actual_fl) // 10  # FL351 vs FL291 = 6 levels, not 60
    diff_ft = abs(plan_alt - actual_alt)

    direction = "高于" if actual_alt > plan_alt else "低于"
    plan_fmt = format_altitude(plan_alt, region)
    actual_fmt = format_altitude(actual_alt, region)

    if diff_levels >= 1:
        return f"{direction}计划高度{diff_levels}个高度层({plan_fmt}→{actual_fmt})"
    else:
        return f"{direction}计划高度约{diff_ft:.0f}ft({plan_fmt}→{actual_fmt})"

# ─── Chart Generation for Appendix (simple Excel-style) ──────────────────────

def generate_profile_chart(plan_wp, actual_pts, dev_results, flight, date_str, output_path):
    """Generate a clean Excel-style plan vs actual altitude profile chart."""
    fig, ax = plt.subplots(figsize=(10, 4.5))

    plan_dists = [w["dist"] for w in plan_wp]
    plan_alts = [w["alt"] for w in plan_wp]
    actual_dists = [r["dis"] for r in dev_results]
    actual_alts = [r["alt"] for r in dev_results]

    ax.plot(plan_dists, plan_alts, '#E8A317', linewidth=2, label='Plan')
    ax.plot(actual_dists, actual_alts, '#1565C0', linewidth=1.5, label='Actual')
    ax.set_xlabel('', fontsize=10)  # distance uncertain, omit label
    ax.set_ylabel('Altitude (ft)', fontsize=10)
    ax.set_title(f'{flight}  {date_str}  Plan vs Actual Altitude Profile', fontsize=12, fontweight='bold')
    ax.legend(loc='upper right', fontsize=9, framealpha=0.8)
    ax.grid(True, alpha=0.3, linewidth=0.5)
    ax.tick_params(labelsize=9)

    plt.tight_layout()
    plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    return output_path

# ─── Word Document Building ──────────────────────────────────────────────────

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color,
    })
    shading.append(shading_elem)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "4472C4")

    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
        if ri % 2 == 1:
            for ci in range(len(headers)):
                set_cell_shading(table.rows[ri + 1].cells[ci], "D6E4F0")
    return table

def generate_main_report(all_results, config, output_path):
    """
    Generate the main analysis report in .docx format.
    all_results: dict keyed by (flight, date) -> analysis_result
    """
    doc = Document()

    # ── Title ──
    title = doc.add_heading('航线优化分析报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    route_name = config.get("route_name", "未指定航线")
    date_range = config.get("date_range", "")

    # ── Analysis Basis ──
    doc.add_heading('分析依据：', level=2)
    total_flights = len(all_results)
    outbound_count = sum(1 for k in all_results if _is_outbound(k[0], config))
    return_count = total_flights - outbound_count

    doc.add_paragraph(
        f'基于{date_range}的实际飞行剖面记录。'
        f'总计提取去程航班{outbound_count}班，回程航班{return_count}班进行交叉比对。'
    )

    # ── Analyze by direction ──
    for direction, dir_label in [("outbound", "去程"), ("return", "回程")]:
        if direction == "outbound":
            dir_results = {k: v for k, v in all_results.items() if _is_outbound(k[0], config)}
        else:
            dir_results = {k: v for k, v in all_results.items() if not _is_outbound(k[0], config)}

        # Get expected flights for this direction
        expected = set()
        if direction == "outbound":
            expected = set(config.get("outbound_flights", []))
        else:
            expected = set(config.get("return_flights", []))

        doc.add_heading(f'{dir_label}航班分析', level=1)

        # Aggregate region statistics across all flights in this direction
        region_deviation_stats = _aggregate_region_stats(dir_results, config)

        if not dir_results:
            doc.add_paragraph('（无数据）')
            continue

        doc.add_paragraph(
            f'在{len(dir_results)}个{dir_label}航班样本中，飞行剖面图反映出实际巡航高度'
            f'在特定空域常与计划高度存在偏差。具体偏差统计如下：'
        )

        # Per-region deviation description
        for rs in region_deviation_stats:
            if rs["deviation_count"] == 0:
                continue
            rate = rs["deviation_count"] / rs["total_flights"] * 100
            dir_text = "偏高" if rs["avg_dev"] > 0 else "偏低"
            plan_fmt = format_altitude(rs.get("plan_alt_ft", 0), rs["region"])
            actual_fmt = format_altitude(rs.get("actual_alt_ft", 0), rs["region"])
            doc.add_paragraph(
                f'{rs["region"]}区域高度{dir_text}（{plan_fmt}→{actual_fmt}）：'
                f'在{rs["total_flights"]}班中出现{rs["deviation_count"]}次，'
                f'占比{rate:.1f}%。'
            )

        # ── Domestic descent analysis ──
        descent_flights = []
        for (flight, date_str), result in dir_results.items():
            da = result.get("descent_analysis")
            if da and da.get("is_premature"):
                descent_flights.append((flight, date_str, da))

        if descent_flights:
            doc.add_heading('下降剖面特征：', level=2)
            # Group by destination
            by_dest = defaultdict(list)
            for flight, date_str, da in descent_flights:
                # The plan TOD waypoint name may just be "TOD", use the region + arrival info instead
                near_wps = da.get("between_waypoints", [])
                dest_key = near_wps[-1] if near_wps else da.get("plan_tod_wp", "?")
                by_dest[dest_key].append((flight, date_str, da))
            for dest_key, items in by_dest.items():
                wp_counts = defaultdict(int)
                for _, _, da in items:
                    wp = da.get("actual_descent_start_wp", "?")
                    wp_counts[wp] += 1
                most_common_wp = max(wp_counts, key=wp_counts.get)
                plan_tod = items[0][2].get("plan_tod_wp", "?")
                diff_nm = items[0][2].get("descent_diff_nm", 0)
                between = items[0][2].get("between_waypoints", [])
                step_suggestion = f'{most_common_wp}点至{plan_tod}点'

                doc.add_paragraph(
                    f'该方向航班均在{most_common_wp}点附近开始下降高度'
                    f'（计划下降点为{plan_tod}，实际提前约{diff_nm:.0f}nm），'
                    f'共{len(items)}班中出现{len(items)}次。'
                )
                if between:
                    doc.add_paragraph(
                        f'建议在{step_suggestion}段设置阶梯下降限制，'
                        f'途经航路点：{"→".join(between[:5])}，'
                        f'避免因提前下降产生额外燃油消耗。'
                    )
            if not descent_flights:
                doc.add_paragraph(
                    '建议对该下降段设置阶梯下降限制，避免因提前下降产生额外燃油消耗。'
                )

        # ── Conclusions ──
        doc.add_heading(f'结论与优化建议（{dir_label}）：', level=2)

        high_rate_regions = [rs for rs in region_deviation_stats
                             if rs["deviation_count"] / max(rs["total_flights"], 1) * 100 >= 50]

        if high_rate_regions:
            for rs in high_rate_regions:
                rate = rs["deviation_count"] / rs["total_flights"] * 100
                actual_fmt = format_altitude(rs.get("actual_alt_ft", 0), rs["region"])
                doc.add_paragraph(
                    f'{rs["region"]}区域的{dir_label}实际巡航高度（低于计划高度，常为{actual_fmt}）'
                    f'出现频次已达到{rate:.1f}%，建议优化高度层。'
                )

            doc.add_heading('优化动作：', level=3)
            for rs in high_rate_regions:
                region = rs["region"]
                plan_fmt = format_altitude(rs.get("plan_alt_ft", 0), region)
                actual_fmt = format_altitude(rs.get("actual_alt_ft", 0), region)
                doc.add_paragraph(
                    f'建议在飞行计划系统（CFP）中，针对{dir_label}的{region}空域航段，'
                    f'人为将计划巡航高度从{plan_fmt}下调至{actual_fmt}，'
                    f'以使计划油量与实际消耗更贴合，避免因长期达不到计划高度而产生额外燃油消耗。'
                )

            low_rate = [rs for rs in region_deviation_stats if rs not in high_rate_regions]
            if low_rate:
                names = "、".join(rs["region"] for rs in low_rate)
                doc.add_paragraph(f'{names}区域的偏差未超过50%，维持现有高度剖面不变。')
        else:
            doc.add_paragraph('所有空域的高度偏差发生率均在可接受范围内。')
            doc.add_paragraph('优化动作：国际段计划巡航高度无需修改。')

        # ── Descent analysis for return flights ──
        if direction == "return" and "return_split" in config:
            doc.add_heading('国内下降剖面特征：', level=3)
            for flight_num, split_info in config["return_split"].items():
                doc.add_paragraph(f'飞往{split_info["dest"]}的航班（{flight_num}）'
                                  f'均在{split_info["descent_keypoint"]}点附近开始下降高度；')

    # ── Summary Table ──
    doc.add_heading('统计表：', level=1)

    headers = ['航班日期', '航班号', '机号', '始发站', '到达站', '备注']
    rows = []
    for (flight, date_str), result in sorted(all_results.items()):
        plan_wp = result.get("_plan_waypoints_raw", [])
        dep = plan_wp[0]["name"] if plan_wp else "?"
        arr = plan_wp[-1]["name"] if plan_wp else "?"
        # Build remarks from warnings
        remarks_parts = []
        for w in result.get("warnings", [])[:3]:
            fl_desc = format_fl_deviation(w.get("plan_alt", 0), w.get("actual_alt", 0), w["region"])
            remarks_parts.append(f'{w["region"]}区域{fl_desc}')
        remarks = "，".join(remarks_parts)
        rows.append([date_str, flight, "", dep, arr, remarks])

    add_styled_table(doc, headers, rows)
    doc.save(output_path)
    print(f"  主报告已生成: {output_path}")
    return output_path


# ─── Appendix Generation ─────────────────────────────────────────────────────

def generate_appendix(all_results, config, output_path, country_index):
    """Generate appendix .docx with per-flight profile charts."""
    doc = Document()
    doc.add_heading('附件：飞行剖面对比图', level=0)

    direction_order = []
    for (flight, date_str), result in sorted(all_results.items()):
        if _is_outbound(flight, config):
            direction_order.append(("去程", flight, date_str, result))
        else:
            direction_order.append(("回程", flight, date_str, result))

    for dir_label, flight, date_str, result in direction_order:
        doc.add_heading(f'{dir_label} — {date_str} {flight}', level=2)

        # Generate chart
        chart_dir = os.path.join(REPORTS_DIR, "charts")
        os.makedirs(chart_dir, exist_ok=True)
        chart_path = os.path.join(chart_dir, f"{flight}_{date_str}_profile.png")

        # Need plan and actual data for chart
        plan_file, actual_file = find_csv_files(flight, date_str)
        if plan_file and actual_file:
            plan_wp = load_plan_track(plan_file)
            actual_pts = load_actual_track(actual_file)

            # Quick deviation recompute for chart
            from analyze_flight import compute_deviations
            dev_results = compute_deviations(actual_pts, plan_wp, {}, country_index)
            from analyze_flight import mark_cruise_points
            mark_cruise_points(dev_results, plan_wp)

            generate_profile_chart(plan_wp, actual_pts, dev_results, flight, date_str, chart_path)

            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(5.5))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add a brief note about findings
        warnings = result.get("warnings", [])
        if warnings:
            notes = []
            for w in warnings[:3]:
                dir_text = "偏高" if w["direction"] == "偏高" else "偏低"
                notes.append(
                    f'{w["region"]}区域{dir_text}：计划FL{int(w["plan_alt"]/100)} '
                    f'实际FL{int(w["actual_alt"]/100)}，偏差{w["avg_dev_ft"]:.0f}ft，持续{w["duration_nm"]:.0f}nm'
                )
            doc.add_paragraph('\n'.join(notes))

    doc.save(output_path)
    print(f"  附件已生成: {output_path}")
    return output_path


# ─── Helpers ─────────────────────────────────────────────────────────────────

def _is_outbound(flight, config):
    """Check if a flight is outbound based on config."""
    outbound_flights = set(config.get("outbound_flights", []))
    return flight in outbound_flights


def _aggregate_region_stats(dir_results, config):
    """
    Aggregate deviation statistics across all flights in one direction,
    grouped by region.
    """
    region_data = defaultdict(lambda: {
        "total_flights": 0,
        "deviation_count": 0,
        "dev_count_total": 0,
        "plan_alts": [],
        "actual_alts": [],
        "devs_signed": [],  # Track signed deviations to determine direction
    })

    for (flight, date_str), result in dir_results.items():
        seen_regions = set()
        for w in result.get("warnings", []):
            reg = w["region"]
            data = region_data[reg]
            data["dev_count_total"] += 1
            data["devs_signed"].append(w["avg_dev_ft"])
            data["plan_alts"].append(w["plan_alt"])
            data["actual_alts"].append(w["actual_alt"])
            if reg not in seen_regions:
                seen_regions.add(reg)
                data["deviation_count"] += 1

        # Count total flights per region
        for reg in set(r["region"] for r in result.get("deviation_data", [])):
            region_data[reg]["total_flights"] += 1

    stats = []
    for reg, data in sorted(region_data.items()):
        if data["total_flights"] == 0:
            continue
        plan_alts_ft = data["plan_alts"]
        actual_alts_ft = data["actual_alts"]
        med_dev = sorted(data["devs_signed"])[len(data["devs_signed"]) // 2] if data["devs_signed"] else 0
        med_plan = sorted(plan_alts_ft)[len(plan_alts_ft) // 2] if plan_alts_ft else 0
        med_actual = sorted(actual_alts_ft)[len(actual_alts_ft) // 2] if actual_alts_ft else 0
        stats.append({
            "region": reg,
            "total_flights": data["total_flights"],
            "deviation_count": data["deviation_count"],
            "avg_dev": med_dev,  # Signed: negative = lower than plan
            "plan_alt_ft": med_plan,
            "actual_alt_ft": med_actual,
            "deviation_rate": data["deviation_count"] / max(data["total_flights"], 1) * 100,
        })
    return stats


# ─── Main ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="生成航班偏差分析 Word 报告")
    parser.add_argument("--flight", "-f", required=True, help="航班号")
    parser.add_argument("--dates", required=True, help="日期列表，逗号分隔 (如 2026-04-02,2026-04-06)")
    parser.add_argument("--route", default=None, help="航线名 (如 鄂州-布鲁塞尔)")
    parser.add_argument("--min-alt-dev", type=float, default=DEFAULT_MIN_ALT_DEV_FT)
    parser.add_argument("--min-dur", type=float, default=DEFAULT_MIN_DURATION_NM)
    parser.add_argument("--output-dir", "-o", default=REPORTS_DIR)
    args = parser.parse_args()

    flight_list = [f.strip() for f in args.flight.split(",")]
    dates = [d.strip() for d in args.dates.split(",")]

    config = {
        "min_alt_deviation_ft": args.min_alt_dev,
        "min_duration_nm": args.min_dur,
        "outbound_flights": flight_list,
        "return_flights": flight_list,
        "route_name": args.route or f"{flight_list[0]}航线",
        "date_range": f"{dates[0]}至{dates[-1]}",
    }

    print("加载国界数据...")
    country_index = CountryIndex(GEOJSON_PATH)

    # Analyze all flights
    all_results = {}
    for flight_num in flight_list:
        for date_str in dates:
            print(f"\n分析: {flight_num} {date_str}")
            plan_file, actual_file = find_csv_files(flight_num, date_str)
            if not plan_file or not actual_file:
                print(f"  SKIP: 找不到 CSV")
                continue
            result = analyze_flight(plan_file, actual_file, config, country_index)
            if result:
                # Also store raw plan waypoints for chart generation
                plan_wp = load_plan_track(plan_file)
                result["_plan_waypoints_raw"] = plan_wp
                all_results[(flight_num, date_str)] = result

    if not all_results:
        print("ERROR: 无有效分析结果")
        sys.exit(1)

    print(f"\n共分析 {len(all_results)} 个航班记录")

    # Generate reports
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    route_name = config["route_name"].replace("/", "-").replace("、", "-")

    main_path = os.path.join(args.output_dir, f"{route_name}_分析报告_{timestamp}.docx")
    appendix_path = os.path.join(args.output_dir, f"{route_name}_附件_{timestamp}.docx")

    generate_main_report(all_results, config, main_path)
    generate_appendix(all_results, config, appendix_path, country_index)

    print(f"\n{'='*60}")
    print("报告生成完毕！")
    print(f"  主报告: {main_path}")
    print(f"  附件:   {appendix_path}")


if __name__ == "__main__":
    main()
