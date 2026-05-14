"""Deterministic report generation for flight deviation analysis.

This module turns the analysis results into deterministic local reports:
- a Word statistics table,
- a chart-only Word appendix,
- the existing integrated HTML report layout with deterministic report text.
"""
import html
import json
import os
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from analyze_flight import (
    BASE_DIR,
    CSV_DIR,
    GEOJSON_PATH,
    RESOURCE_DIR,
    REPORTS_DIR,
    CountryIndex,
    compute_deviations,
    find_csv_files,
    load_actual_track,
    load_plan_track,
    mark_cruise_points,
    haversine_distance,
)
from generate_docx_report import generate_profile_chart, nearest_valid_fl
from flight_keys import display_datetime_from_key, parse_flight_key


METADATA_PATH = os.path.join(CSV_DIR, "metadata.json")

AIRPORT_NAMES = {
    "ZHEC": "鄂州",
    "EBBR": "布鲁塞尔",
    "ZSFZ": "福州",
    "ZWWW": "乌鲁木齐",
    "ZBAA": "北京",
    "ZSPD": "上海浦东",
    "ZGGG": "广州",
    "ZGSZ": "深圳",
    "ZUCK": "重庆",
    "ZUUU": "成都",
    "ZBTJ": "天津",
    "ZHCC": "郑州",
    "VHHH": "香港",
}

COUNTRY_NAMES = {
    "China": "中国",
    "Mongolia": "蒙古",
    "Russia": "俄罗斯",
    "Kazakhstan": "哈萨克斯坦",
    "Norway": "挪威",
    "Sweden": "瑞典",
    "Finland": "芬兰",
    "Denmark": "丹麦",
    "Germany": "德国",
    "Netherlands": "荷兰",
    "Belgium": "比利时",
    "Poland": "波兰",
    "Belarus": "白俄罗斯",
    "Ukraine": "乌克兰",
    "United Kingdom": "英国",
    "Ireland": "爱尔兰",
    "Turkey": "土耳其",
    "Latvia": "拉脱维亚",
    "Lithuania": "立陶宛",
    "Estonia": "爱沙尼亚",
    "Czechia": "捷克",
    "Czech Republic": "捷克",
    "Slovakia": "斯洛伐克",
    "Austria": "奥地利",
    "Hungary": "匈牙利",
    "Romania": "罗马尼亚",
    "Bulgaria": "保加利亚",
    "Georgia": "格鲁吉亚",
    "Azerbaijan": "阿塞拜疆",
    "Armenia": "亚美尼亚",
    "Uzbekistan": "乌兹别克斯坦",
    "Kyrgyzstan": "吉尔吉斯斯坦",
    "Tajikistan": "塔吉克斯坦",
    "Turkmenistan": "土库曼斯坦",
    "Pakistan": "巴基斯坦",
    "India": "印度",
}

CHINESE_NUMBERS = {
    1: "一个",
    2: "两个",
    3: "三个",
    4: "四个",
    5: "五个",
    6: "六个",
}

SEA_LABELS = {"XX", "海上", "海上/公海"}
EUROPE_LABELS = {"欧洲", "Europe"}
DOMESTIC_LABELS = {"国内段", "中国", "CN"}


def load_metadata():
    if not os.path.exists(METADATA_PATH):
        return {}
    try:
        with open(METADATA_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except (OSError, json.JSONDecodeError):
        return {}


def save_metadata_entry(key, data):
    os.makedirs(CSV_DIR, exist_ok=True)
    metadata = load_metadata()
    metadata[key] = {k: v for k, v in data.items() if v is not None}
    with open(METADATA_PATH, "w", encoding="utf-8") as f:
        json.dump(metadata, f, ensure_ascii=False, indent=2)


def generate_local_reports(analysis, config, output_dir=REPORTS_DIR):
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    route_name = _route_name(analysis)
    safe_name = _safe_filename(route_name)

    stats_path = os.path.join(output_dir, f"{safe_name}_stats_{ts}.docx")
    appendix_path = os.path.join(output_dir, f"{safe_name}_appendix_{ts}.docx")
    html_path = os.path.join(output_dir, f"{safe_name}_integrated_{ts}.html")

    rows = build_statistics_rows(analysis)
    generate_statistics_docx(rows, stats_path)
    generate_chart_appendix(analysis, appendix_path)
    generate_integrated_html_local(analysis, rows, config, html_path)

    return {
        "stats": stats_path,
        "appendix": appendix_path,
        "html": html_path,
    }


def generate_stats_report(analysis, output_dir=REPORTS_DIR):
    """Generate only the Word statistics table."""
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    safe_name = _safe_filename(_route_name(analysis))
    stats_path = os.path.join(output_dir, f"{safe_name}_stats_{ts}.docx")
    rows = build_statistics_rows(analysis)
    return generate_statistics_docx(rows, stats_path)


def generate_appendix_report(analysis, output_dir=REPORTS_DIR):
    """Generate only the chart appendix Word report."""
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    safe_name = _safe_filename(_route_name(analysis))
    appendix_path = os.path.join(output_dir, f"{safe_name}_appendix_{ts}.docx")
    return generate_chart_appendix(analysis, appendix_path)


def generate_table_html_report(analysis, config, output_dir=REPORTS_DIR):
    """Generate the integrated HTML report with deterministic table text only."""
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    safe_name = _safe_filename(_route_name(analysis))
    html_path = os.path.join(output_dir, f"{safe_name}_integrated_{ts}.html")
    rows = build_statistics_rows(analysis)
    return generate_integrated_html_local(analysis, rows, config, html_path)


def build_statistics_rows(analysis):
    metadata = load_metadata()
    rows = []
    for key, result in sorted(analysis.get("flights", {}).items(), key=_flight_item_sort_key):
        flight, date_str = _split_key(key)
        meta = metadata.get(key, {})
        result_meta = result.get("metadata", {})
        dep = _display_airport(meta.get("dep_airport") or result_meta.get("dep_airport"))
        arr = _display_airport(meta.get("arr_airport") or result_meta.get("arr_airport"))
        aircraft = meta.get("aircraft", "")
        date_text = _format_date_time(date_str, meta.get("tko_time") or result_meta.get("dep_time"))
        rows.append({
            "date": date_text,
            "flight": flight,
            "aircraft": _normalize_aircraft(aircraft),
            "dep": dep,
            "arr": arr,
            "remarks": build_remarks(result),
        })
    return rows


def build_remarks(result):
    parts = []
    seen = set()
    for warning in result.get("warnings", []):
        region_phrase = _display_region_phrase(result, warning)
        actual_alt = warning.get("actual_alt", 0)
        plan_alt = warning.get("plan_alt", 0)
        direction = "低于" if warning.get("avg_dev_ft", 0) < 0 else "高于"
        actual_fl = _format_fl(actual_alt, result, warning)
        level_count = _height_level_count(actual_alt, plan_alt)
        text = f"{region_phrase}高度 {actual_fl} {direction}计划高度{level_count}高度层"
        if text not in seen:
            seen.add(text)
            parts.append(text)
        if len(parts) >= 4:
            break
    descent = result.get("descent_analysis") or {}
    if descent.get("is_premature") and _is_domestic_descent(descent) and len(parts) < 4:
        wp = str(descent.get("actual_descent_start_wp") or "").strip()
        dist = float(descent.get("descent_diff_nm") or 0)
        context = f"{wp}点附近" if wp else "实际下降点附近"
        text = f"国内段（{context}）较计划提前下降约{dist:.0f}nm"
        if text not in seen:
            parts.append(text)
    return "，\n".join(parts)


def generate_statistics_docx(rows, output_path):
    doc = Document()
    _setup_page(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("统计表：")
    _set_run_font(run, Pt(16), bold=True)

    headers = ["航班日期", "航班号", "机号", "始发站", "到达站", "备注"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        _shade_cell(cell, "CFE8CC")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        _set_run_font(r, Pt(12), bold=True)

    for row in rows:
        cells = table.add_row().cells
        values = [row["date"], row["flight"], row["aircraft"], row["dep"], row["arr"], row["remarks"]]
        for idx, value in enumerate(values):
            cells[idx].text = ""
            _shade_cell(cells[idx], "CFE8CC")
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 5 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            _set_run_font(r, Pt(11))

    widths = [Cm(3.4), Cm(2.2), Cm(1.8), Cm(2.5), Cm(2.8), Cm(9.2)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_chart_appendix(analysis, output_path):
    doc = Document()
    _setup_page(doc)
    chart_dir = os.path.join(REPORTS_DIR, "charts")
    os.makedirs(chart_dir, exist_ok=True)
    country_index = CountryIndex(GEOJSON_PATH)

    for key, _result in sorted(analysis.get("flights", {}).items(), key=_flight_item_sort_key):
        flight, date_str = _split_key(key)
        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run(f"{date_str}  {flight}")
        _set_run_font(run, Pt(12), bold=True)

        plan_file, actual_file = find_csv_files(flight, date_str)
        if plan_file and actual_file:
            plan_wp = load_plan_track(plan_file)
            actual_pts = load_actual_track(actual_file)
            dev_results = compute_deviations(actual_pts, plan_wp, {}, country_index)
            mark_cruise_points(dev_results, plan_wp)
            chart_path = os.path.join(chart_dir, f"{key}_profile.png")
            generate_profile_chart(plan_wp, actual_pts, dev_results, flight, date_str, chart_path)
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(6.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


def generate_integrated_html_local(analysis, rows, config, output_path):
    external_template = os.path.join(BASE_DIR, "report_integrated.html")
    bundled_template = os.path.join(RESOURCE_DIR, "report_integrated.html")
    template_path = external_template if os.path.exists(external_template) else bundled_template
    with open(template_path, "r", encoding="utf-8") as f:
        template = f.read()

    all_days = {}
    for key, result in sorted(analysis.get("flights", {}).items(), key=_flight_item_sort_key):
        slim_dev = []
        for d in result.get("deviation_data", []):
            slim_dev.append({
                "t": d.get("t", ""),
                "lat": d.get("lat", 0),
                "lon": d.get("lon", 0),
                "alt": d.get("alt", 0),
                "dis": d.get("dis", 0),
                "plan_alt": d.get("plan_alt", 0),
                "plan_lat": d.get("plan_lat", 0),
                "plan_lon": d.get("plan_lon", 0),
                "xt": d.get("xt", d.get("cross_track_nm", 0)),
                "alt_dev": d.get("alt_dev", d.get("alt_dev_ft", 0)),
                "fuel_dev": d.get("fuel_dev", d.get("fuel_dev_lbs", 0)),
                "region": d.get("region", ""),
                "country": d.get("country", ""),
                "country_label": _translate_country(str(d.get("country", "") or "")),
                "is_cruise": d.get("is_cruise", False),
            })
        all_days[key] = {
            "metadata": result.get("metadata", {}),
            "plan_waypoints": result.get("plan_waypoints", []),
            "actual_track": result.get("actual_track", []),
            "diagnostics": result.get("diagnostics", {}),
            "deviation_data": slim_dev,
            "warnings": result.get("warnings", []),
            "descent_analysis": result.get("descent_analysis"),
        }

    route_name = _route_name(analysis)
    report_html = statistics_rows_to_html(rows)
    html_doc = template.replace("{{TITLE}}", f"{route_name} 偏差分析报告")
    html_doc = html_doc.replace("{{FLIGHT}}", route_name)
    html_doc = html_doc.replace("{{H_THRESHOLD}}", str(config.get("min_alt_deviation_ft", 1000)))
    html_doc = html_doc.replace("{{V_THRESHOLD}}", str(config.get("min_duration_nm", 50)))
    html_doc = html_doc.replace("{{ALL_DAYS_JSON}}", json.dumps(all_days, ensure_ascii=False))
    html_doc = html_doc.replace("{{REPORT_MD_JSON}}", json.dumps("", ensure_ascii=False))
    html_doc = html_doc.replace("{{REPORT_TEXT}}", report_html)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_doc)
    return output_path


def statistics_rows_to_html(rows):
    headers = ["航班日期", "航班号", "机号", "始发站", "到达站", "备注"]
    body = ["<h2>统计表</h2>", "<table>"]
    body.append("<tr>" + "".join(f"<th>{html.escape(h)}</th>" for h in headers) + "</tr>")
    for row in rows:
        values = [row["date"], row["flight"], row["aircraft"], row["dep"], row["arr"], row["remarks"]]
        body.append("<tr>" + "".join(f"<td>{html.escape(v).replace(chr(10), '<br>')}</td>" for v in values) + "</tr>")
    body.append("</table>")
    return "\n".join(body)


def _display_region_phrase(result, warning):
    region = str(warning.get("region", "") or "")
    countries = _country_labels_for_warning(result, warning)

    if _is_domestic_region(region, countries):
        context = _waypoint_context(result, warning)
        return f"国内段（{context}）" if context else "国内段"

    if region in EUROPE_LABELS:
        if len(countries) > 1:
            return f"欧洲区域（{'、'.join(countries[:5])}）"
        if len(countries) == 1:
            return f"{countries[0]}区域"
        return "欧洲区域"

    if region not in ("海上", "海上/公海"):
        return f"{region}区域"

    nearest = _nearest_country_name(result, warning)
    return f"{nearest}附近区域" if nearest else "近海区域"


def _display_region(result, warning):
    phrase = _display_region_phrase(result, warning)
    return phrase[:-2] if phrase.endswith("区域") else phrase


def _is_domestic_descent(descent):
    if "is_domestic_descent" in descent:
        return bool(descent.get("is_domestic_descent"))
    if descent.get("actual_descent_start_iso") == "CN":
        return True
    return str(descent.get("region", "") or "") in DOMESTIC_LABELS


def _country_labels_for_warning(result, warning):
    raw = list(warning.get("countries") or [])
    if not raw:
        raw = [
            str(p.get("country", "") or "")
            for p in _warning_event_points(result, warning)
        ]

    labels = []
    for country in raw:
        if not country or country in SEA_LABELS:
            continue
        label = _translate_country(country)
        if label not in labels:
            labels.append(label)
    return labels


def _warning_event_points(result, warning):
    data = result.get("deviation_data", [])
    start = float(warning.get("start_dist", warning.get("start_dis", 0)) or 0)
    end = float(warning.get("end_dist", warning.get("end_dis", start)) or start)
    if end < start:
        start, end = end, start
    return [
        p for p in data
        if start <= float(p.get("dis", 0) or 0) <= end
    ]


def _is_domestic_region(region, countries):
    if region in DOMESTIC_LABELS:
        return True
    return any(country == "中国" for country in countries)


def _waypoint_context(result, warning):
    plan_wps = result.get("plan_waypoints", [])
    if not plan_wps:
        return ""

    points = _warning_event_points(result, warning)
    if points:
        target = max(points, key=lambda p: abs(float(p.get("alt_dev", 0) or 0)))
        dist = float(target.get("dis", 0) or 0)
    else:
        start = float(warning.get("start_dist", warning.get("start_dis", 0)) or 0)
        end = float(warning.get("end_dist", warning.get("end_dis", start)) or start)
        dist = (start + end) / 2

    nearest = min(plan_wps, key=lambda w: abs(float(w.get("dist", 0) or 0) - dist))
    nearest_delta = abs(float(nearest.get("dist", 0) or 0) - dist)
    if nearest.get("name") and nearest_delta <= 20:
        return f"{nearest['name']}点附近"

    prev_wp = None
    next_wp = None
    for wp in plan_wps:
        wp_dist = float(wp.get("dist", 0) or 0)
        if wp_dist <= dist:
            prev_wp = wp
        elif wp_dist > dist:
            next_wp = wp
            break
    if prev_wp and next_wp and prev_wp.get("name") and next_wp.get("name"):
        return f"{prev_wp['name']}-{next_wp['name']}航段"
    if nearest.get("name"):
        return f"{nearest['name']}点附近"
    return ""


def _nearest_country_name(result, warning):
    data = result.get("deviation_data", [])
    start = float(warning.get("start_dist", warning.get("start_dis", 0)) or 0)
    end = float(warning.get("end_dist", warning.get("end_dis", start)) or start)
    event_points = [p for p in data if start <= float(p.get("dis", 0) or 0) <= end]
    if not event_points:
        event_points = data
    if not event_points:
        return ""

    mid = event_points[len(event_points) // 2]
    lat = float(mid.get("lat", 0) or 0)
    lon = float(mid.get("lon", 0) or 0)

    candidates = []
    for point in data:
        country = str(point.get("country", "") or "")
        region = str(point.get("region", "") or "")
        if not country or country in ("XX", "海上/公海", "海上") or region in ("海上", "海上/公海"):
            continue
        candidates.append(point)
    if not candidates:
        return ""

    nearest = min(
        candidates,
        key=lambda p: haversine_distance(lat, lon, float(p.get("lat", 0) or 0), float(p.get("lon", 0) or 0)),
    )
    return _translate_country(str(nearest.get("country", "") or ""))


def _translate_country(name):
    return COUNTRY_NAMES.get(name, name)


def _split_key(key):
    parsed = parse_flight_key(key)
    return parsed.flight, parsed.date_key


def _flight_item_sort_key(item):
    parsed = parse_flight_key(item[0])
    return (parsed.date, parsed.time_token, parsed.suffix, parsed.flight)


def _route_name(analysis):
    flights = sorted({_split_key(key)[0] for key in analysis.get("flights", {})})
    return "/".join(flights) if flights else "航班分析"


def _safe_filename(name):
    safe = "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in name)
    return safe.strip("_") or "flight_report"


def _display_airport(value):
    if not value:
        return ""
    value = str(value).strip()
    return AIRPORT_NAMES.get(value, value)


def _normalize_aircraft(value):
    value = str(value or "").strip()
    if value.upper().startswith("B-"):
        value = value[2:]
    return value


def _format_date_time(date_str, time_str):
    parsed = parse_flight_key(f"X_{date_str}")
    if parsed.time_token:
        return display_datetime_from_key(f"X_{date_str}")
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%Y/%m/%d")
    except ValueError:
        return date_str


def _format_fl(alt_ft, result=None, warning=None):
    region = str((warning or {}).get("region", "") or "")
    countries = _country_labels_for_warning(result or {}, warning or {}) if warning else []
    if _is_domestic_region(region, countries):
        fl = nearest_valid_fl(float(alt_ft or 0))
    else:
        fl = _standard_icao_fl(alt_ft)
    return f"FL{fl:03d}"


def _standard_icao_fl(alt_ft):
    """Round non-China altitudes to ordinary ICAO whole-thousand flight levels."""
    return max(0, int(round(float(alt_ft or 0) / 1000.0)) * 10)


def _height_level_count(actual_alt, plan_alt):
    diff = abs(float(actual_alt or 0) - float(plan_alt or 0))
    levels = max(1, int(round(diff / 2000.0)))
    return CHINESE_NUMBERS.get(levels, f"{levels}个")


def _setup_page(doc):
    for section in doc.sections:
        section.page_width = Cm(29.7)
        section.page_height = Cm(21)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def _set_run_font(run, size, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = size
    run.bold = bold


def _shade_cell(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)
