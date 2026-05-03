"""
AI-Powered Flight Deviation Report Generator
=============================================
Uses DeepSeek API to generate natural-language analysis reports.

Usage:
    # Full pipeline (analysis + AI draft + docx)
    python review_report.py --flight I99806 --dates ... --reference "原始报告.docx"

    # Prompt only (review before API call)
    python review_report.py --data analysis.json --prompt-only

    # Finalize approved draft to docx
    python review_report.py --data analysis.json --draft draft.md --finalize
"""
import os, sys, json, argparse, textwrap
from datetime import datetime
from collections import defaultdict
import requests

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from analyze_flight import (
    analyze_flight, find_csv_files, CountryIndex, GEOJSON_PATH, CSV_DIR,
    DEFAULT_MIN_ALT_DEV_FT, DEFAULT_MIN_DURATION_NM, M_TO_FT,
    load_plan_track, load_actual_track, compute_deviations, mark_cruise_points,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORTS_DIR = os.path.join(BASE_DIR, "reports")
CONFIG_PATH = os.path.join(BASE_DIR, "deepseek_config.json")

# ─── Region Normalization ─────────────────────────────────────────────────────

REGION_NORMALIZE = {
    "国内段": "国内段", "蒙古": "蒙古", "俄罗斯": "俄罗斯",
    "哈萨克斯坦": "哈萨克斯坦", "海上": "海上", "欧洲": "欧洲",
    "阿塞拜疆": "欧洲", "格鲁吉亚": "欧洲", "土耳其": "欧洲",
    "希腊": "欧洲", "波兰": "欧洲", "德国": "欧洲", "比利时": "欧洲",
    "白俄罗斯": "欧洲", "乌克兰": "欧洲", "罗马尼亚": "欧洲", "保加利亚": "欧洲",
    "土库曼斯坦": "中亚", "乌兹别克斯坦": "中亚",
    # ISO-based codes that slip through
    "其他(AZ)": "欧洲", "其他(AM)": "欧洲", "其他(GE)": "欧洲",
}

def normalize_region(r):
    if r in REGION_NORMALIZE: return REGION_NORMALIZE[r]
    if r.startswith("其他("): r = r[3:-1]
    return REGION_NORMALIZE.get(r, r)


# ─── DeepSeek API ────────────────────────────────────────────────────────────

def load_api_config():
    if not os.path.exists(CONFIG_PATH):
        print(f"ERROR: {CONFIG_PATH} not found"); sys.exit(1)
    with open(CONFIG_PATH) as f: return json.load(f)

def call_deepseek(messages, config):
    url = config["base_url"] + "/chat/completions"
    resp = requests.post(url,
        headers={"Authorization": f"Bearer {config['api_key']}", "Content-Type": "application/json"},
        json={"model": config["model"], "messages": messages,
              "max_tokens": config.get("max_tokens", 4096),
              "temperature": config.get("temperature", 0.3)})
    resp.raise_for_status()
    return resp.json()["choices"][0]["message"]["content"]


# ─── Prompt Builder ──────────────────────────────────────────────────────────

def build_system_prompt(ref_text="", appendix_text=""):
    base = textwrap.dedent("""\
    你是一名资深航班运行分析专家。请根据飞行偏差数据生成航线优化分析报告。

    ## 报告结构（必须严格遵循）
    报告由以下章节组成，每个章节的标题和内容范围如下：

    ### 分析依据
    一段话。说明数据来源（日期范围、航班号、航线、样本量、去程/回程）。

    ### [方向]航班分析（如"去程航班分析（鄂州 - 布鲁塞尔，I98831/I98833）"）
    先一句话总述。然后按区域逐个列出偏差统计，每个区域一行：
    "[区域]区域高度[偏高/偏低]（[FLxxx/xxxx米]）：在[N]班中出现[X]次，占比 [XX.X]%。"
    - 只列发生率 >= 50% 的区域
    - 国际段用 FL 格式（FL301, FL310, FL320...）
    - 国内段用中国米制标准（如 10700米(FL351)）
    - 如果该区域内偏差复杂（有高有低），写"[区域]区域高度异常（FLxxx/FLxxx等）：..."

    ### 结论与优化建议（[方向]）
    先一句话总结。然后"优化动作："具体列出 CFP 修改建议：
    - 每个建议：[动作描述] + [理由]
    - 发生率 < 50% 的区域统一写"偏差未超过50%，维持现有高度剖面不变"
    - 国内段提前下降需给出阶梯下降方案

    ### 国内下降剖面特征（如有提前下降）
    "飞往[目的地]的航班（[航班号]）均在[航路点]点附近开始下降高度；"
    一句话描述即可。然后给阶梯下降建议。

    ### 统计表
    Markdown 表格，列出：日期、航班号、始发站、到达站、备注（用自然语言描述偏差）

    ## 写作规则
    - **只使用用户提供的数据中明确列出的日期，绝对不要编造日期**
    - 统计表中的每一行日期必须与用户提供的日期列表完全一致
    - 每句话有数字支撑，不空泛
    - 高度层偏差用"X个高度层"描述（不写 FL 编号差值）
    - 国际段 FL 格式：FL290, FL310, FL330...
    - 中国段米制格式：10700米(FL351), 8900米(FL291)...
    - 语言简洁专业，每段一个核心结论
    - 不要展开"典型事件"、"极端案例"等子分析
    - 不要在正文中列出完整表格（表格只在统计表章节出现）
    - 不要给多个备选方案，只给一个明确建议
    - 直接输出报告正文，不要加开场白
    - **必须使用 Markdown 标题标记**：# 用于报告大标题，## 用于章节标题，### 用于子标题
    - 不要用纯文本加粗代替标题
    """)

    if ref_text:
        base += textwrap.dedent(f"""

            ## 参考报告（模仿行文风格和措辞习惯，但不要照搬其航线名称和航班号）
            === 参考报告开始 ===
            {ref_text[:6000]}
            === 参考报告结束 ===
            注意：以上参考报告的航线、航班号与你本次要写的不同。请使用本次数据中的航线名和航班号。
            """)

    if appendix_text:
        base += textwrap.dedent(f"""

            ## 参考附件格式（每个航班一张剖面图 + 简短文字说明）
            === 参考附件开始 ===
            {appendix_text[:3000]}
            === 参考附件结束 ===
            """)

    return base


def build_user_prompt(analysis, config):
    """Build data prompt — only aggregate stats, no per-flight raw data."""

    from generate_docx_report import format_altitude as fmt_alt

    # Step 1: Group flights by direction
    flights_by_dir = {"outbound": [], "return": []}
    for key, result in analysis.get("flights", {}).items():
        parts = key.rsplit("_", 1)
        flight = parts[0]
        date_str = parts[1] if len(parts) > 1 else ""
        flights_by_dir["outbound"].append((flight, date_str, result))

    # Step 2: Per-direction aggregate stats
    dir_sections = []
    for direction, dir_label in [("outbound", "去程"), ("return", "回程")]:
        entries = flights_by_dir[direction]
        if not entries:
            continue

        # Aggregate region deviation rates across all flights
        region_counts = defaultdict(lambda: {"dev": 0, "total": 0, "plan_alts": [], "actual_alts": []})
        for flight, date_str, result in entries:
            seen_regions = set()
            for w in result.get("warnings", []):
                reg = normalize_region(w["region"])
                rdata = region_counts[reg]
                rdata["plan_alts"].append(w["plan_alt"])
                rdata["actual_alts"].append(w["actual_alt"])
                if reg not in seen_regions:
                    seen_regions.add(reg)
                    rdata["dev"] += 1
            for reg in set(normalize_region(r["region"]) for r in result.get("deviation_data", [])):
                region_counts[reg]["total"] += 1

        # Format region stats
        region_lines = []
        for reg in sorted(region_counts):
            rdata = region_counts[reg]
            if rdata["total"] == 0:
                continue
            rate = rdata["dev"] / rdata["total"] * 100
            if rate < 50:
                region_lines.append(f"  {reg}区域：发生率{rate:.1f}%（低于50%，不纳入优化）")
                continue
            median_plan = sorted(rdata["plan_alts"])[len(rdata["plan_alts"])//2] if rdata["plan_alts"] else 0
            median_actual = sorted(rdata["actual_alts"])[len(rdata["actual_alts"])//2] if rdata["actual_alts"] else 0
            plan_fl_fmt = fmt_alt(median_plan, reg)
            actual_fl_fmt = fmt_alt(median_actual, reg)
            region_lines.append(
                f"  {reg}区域：计划{plan_fl_fmt} 实际{actual_fl_fmt}，{rdata['dev']}/{rdata['total']}班，发生率{rate:.1f}%"
            )

        # Descent summary
        descent_summary = ""
        for flight, date_str, result in entries:
            da = result.get("descent_analysis")
            if da and da.get("is_premature"):
                if not descent_summary:
                    descent_summary = (
                        f"  提前下降：实际在{da['actual_descent_start_wp']}点"
                        f"（计划TOD点{da['plan_tod_wp']}，提前约{da['descent_diff_nm']:.0f}nm）"
                    )
                else:
                    descent_summary += f"；另在{da['actual_descent_start_wp']}点"

        dir_sections.append(textwrap.dedent(f"""
            ## {dir_label}航班（{len(entries)}班）
            ### 区域偏差统计
            {chr(10).join(region_lines) if region_lines else '  （无显著偏差）'}
            ### 下降分析
            {descent_summary or '  （未检测到提前下降）'}
        """))

    # Step 3: Per-flight summary for statistics table (1 line per flight)
    flight_summaries = []
    for key, result in sorted(analysis.get("flights", {}).items()):
        parts = key.rsplit("_", 1)
        flight, date_str = parts[0], parts[1] if len(parts) > 1 else ""
        warnings = result.get("warnings", [])
        if warnings:
            by_region = defaultdict(list)
            for w in warnings:
                reg = normalize_region(w["region"])
                by_region[reg].append(w)
            notes = []
            for reg, wlist in sorted(by_region.items()):
                w = wlist[0]
                dir_text = w["direction"]
                plan_fl = fmt_alt(w["plan_alt"], reg)
                actual_fl = fmt_alt(w["actual_alt"], reg)
                notes.append(f"{reg}区域{dir_text}({plan_fl}→{actual_fl})")
            da = result.get("descent_analysis")
            if da and da.get("is_premature"):
                notes.append(f"在{da['actual_descent_start_wp']}点提前下降")
            flight_summaries.append(
                f"  {date_str} | {flight} | 鄂州 | 雅典 | {'，'.join(notes)}"
            )
        else:
            flight_summaries.append(f"  {date_str} | {flight} | 鄂州 | 雅典 | 无明显偏差")

    # Step 4: Domestic waypoint summary (only key waypoints with >500ft dev)
    wp_lines = []
    for direction_entries in [flights_by_dir["outbound"], flights_by_dir["return"]]:
        for flight, date_str, result in direction_entries:
            wp_devs = defaultdict(list)
            for d in result.get("deviation_data", []):
                if d.get("region") == "国内段" and d.get("is_cruise"):
                    plan_wps = result.get("plan_waypoints", [])
                    nearest = min(plan_wps, key=lambda w: abs(w["dist"] - d["dis"])) if plan_wps else None
                    if nearest:
                        wp_devs[nearest["name"]].append(d["alt"] - d["plan_alt"])
            for wp_name, devs in sorted(wp_devs.items()):
                median_dev = sorted(devs)[len(devs)//2]
                if abs(median_dev) > 500:
                    wp_lines.append(f"  {flight}_{date_str} {wp_name}: 偏差{median_dev:.0f}ft")

    # Build explicit date list to prevent AI hallucination
    date_list = sorted(set(
        k.rsplit("_", 1)[1] for k in analysis.get("flights", {})
    ))
    date_str_explicit = "、".join(date_list)

    prompt = textwrap.dedent(f"""
        航线: {config.get('route_name', '未指定')}
        分析日期（仅这些，不要编造其他日期）: {date_str_explicit}
        总航班数: {len(analysis.get('flights', {}))}

        {"".join(dir_sections)}

        ## 统计表数据（每班一行，按此填充统计表，不要编造日期或备注）
        {chr(10).join(flight_summaries)}

        ## 国内段航路点偏差（供正文参考）
        {chr(10).join(wp_lines[:30]) if wp_lines else '（无显著航路点偏差）'}

        请按系统提示词的结构要求生成完整报告，统计表必须使用上面"统计表数据"中的每一行。
        直接输出报告正文，不要加开场白。
    """)
    return prompt


# ─── Docx Finalizer (matching original report formatting) ────────────────────

def generate_integrated_html(analysis, ai_report_md, config, output_path):
    """Generate integrated HTML report embedding all analysis data + AI text."""
    import json as _json
    template_path = os.path.join(BASE_DIR, "report_integrated.html")
    if not os.path.exists(template_path):
        print("  WARNING: report_integrated.html not found, skipping HTML generation")
        return None
    with open(template_path, "r", encoding="utf-8") as f:
        template = f.read()

    # Build per-flight data for JS
    all_days = {}
    for key, result in sorted(analysis.get("flights", {}).items()):
        parts = key.rsplit("_", 1)
        date_str = parts[1] if len(parts) > 1 else key
        all_days[date_str] = {
            "metadata": result.get("metadata", {}),
            "plan_waypoints": result.get("plan_waypoints", []),
            "deviation_data": result.get("deviation_data", []),
            "warnings": result.get("warnings", []),
            "descent_analysis": result.get("descent_analysis"),
        }

    flight_name = "航班分析"
    for key in analysis.get("flights", {}):
        flight_name = key.split("_")[0]
        break

    html = template.replace("{{TITLE}}", f"{flight_name} 偏差分析报告")
    html = html.replace("{{FLIGHT}}", flight_name)
    html = html.replace("{{H_THRESHOLD}}", str(config.get("min_alt_deviation_ft", 1000)))
    html = html.replace("{{V_THRESHOLD}}", str(config.get("min_duration_nm", 200)))
    html = html.replace("{{ALL_DAYS_JSON}}", _json.dumps(all_days, ensure_ascii=False))
    html = html.replace("{{REPORT_MD_JSON}}", _json.dumps(ai_report_md, ensure_ascii=False))

    # Simple markdown → HTML for report text
    md_html = _md_to_html(ai_report_md)
    html = html.replace("{{REPORT_TEXT}}", md_html)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"  整合 HTML 报告已生成: {output_path}")
    return output_path


def _md_to_html(md_text):
    """Basic markdown to HTML converter for report text."""
    lines = md_text.split("\n")
    result = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            result.append("<br>")
            continue
        if stripped.startswith("|") and "|" in stripped[1:]:
            if not in_table:
                result.append('<table class="md-table">')
                in_table = True
            cells = stripped.split("|")[1:-1]
            is_header = all(c.strip().startswith(":") or c.strip().startswith("-") for c in cells if c.strip())
            if is_header:
                continue
            tag = "th" if in_table and not result[-1].startswith("<tr>") else "td"
            result.append("<tr>" + "".join(f"<{tag}>{c.strip()}</{tag}>" for c in cells) + "</tr>")
            continue
        else:
            if in_table:
                result.append("</table>")
                in_table = False
        if stripped.startswith("# ") and not stripped.startswith("## "):
            result.append(f"<h1>{stripped[2:]}</h1>")
        elif stripped.startswith("## "):
            result.append(f"<h2>{stripped[3:]}</h2>")
        elif stripped.startswith("### "):
            result.append(f"<h3>{stripped[4:]}</h3>")
        elif stripped.startswith("- "):
            result.append(f"<p>{stripped}</p>")
        else:
            result.append(f"<p>{stripped}</p>")
    if in_table:
        result.append("</table>")
    return "\n".join(result)


def finalize_to_docx(draft_path, analysis, output_path):
    """Convert markdown draft to .docx matching the original report's formatting:
    - 宋体 throughout
    - Title: 22pt bold centered
    - Section headers: 12pt bold
    - Body: 10.5pt
    - Table: 9pt
    - A4 page, standard margins
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    FONT_NAME = "宋体"
    SIZE_TITLE = Pt(22)
    SIZE_H1 = Pt(14)
    SIZE_H2 = Pt(12)
    SIZE_BODY = Pt(10.5)
    SIZE_TABLE = Pt(9)

    def add_run(para, text, size=SIZE_BODY, bold=False, color=None):
        run = para.add_run(text)
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run.font.size = size
        run.bold = bold
        if color:
            run.font.color.rgb = color

    with open(draft_path, "r", encoding="utf-8") as f:
        markdown = f.read()

    doc = Document()

    # Page setup: A4
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Parse and render content
    lines = markdown.split("\n")
    title_seen = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        if not line:
            continue

        # Title: "# xxx" or standalone first line
        if (line.startswith("# ") and not line.startswith("## ")) or (not title_seen and not line.startswith("#")):
            title_seen = True
            text = line[2:] if line.startswith("# ") else line
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, text, SIZE_TITLE, bold=True)
            doc.add_paragraph()
            continue

        if line.startswith("## "):
            p = doc.add_paragraph()
            add_run(p, line[3:], SIZE_H1, bold=True)

        elif line.startswith("### "):
            p = doc.add_paragraph()
            add_run(p, line[4:], SIZE_H2, bold=True)

        elif line.startswith("|") and "|" in line[1:] and not line.startswith("|-"):
            pass  # tables handled below

        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_run(p, "• " + line[2:])

        else:
            # Regular paragraph
            text = line.replace("**", "").replace("*", "").replace("`", "")
            if text:
                p = doc.add_paragraph()
                add_run(p, text)

    # Process tables
    table_lines = []
    in_table = False
    for line in markdown.split("\n"):
        line = line.strip()
        if line.startswith("|") and not line.startswith("|-"):
            if not in_table:
                if table_lines:
                    _render_table(doc, table_lines, FONT_NAME, SIZE_TABLE)
                table_lines = []
                in_table = True
            table_lines.append(line)
        elif in_table and line.startswith("|-"):
            continue  # skip separator
        else:
            if in_table:
                in_table = False
    if table_lines:
        _render_table(doc, table_lines, FONT_NAME, SIZE_TABLE)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"  最终报告已生成: {output_path}")

    # Generate appendix
    appendix_path = output_path.replace(".docx", "_附件.docx")
    _generate_appendix(analysis, appendix_path, FONT_NAME)
    print(f"  附件已生成: {appendix_path}")
    return output_path


def _render_table(doc, lines, font_name, font_size):
    """Render markdown table lines as a styled Word table."""
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    rows = []
    for line in lines:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
    if not rows:
        return

    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"

    for ri, row_data in enumerate(rows):
        for ci in range(ncols):
            cell = table.rows[ri].cells[ci]
            text = row_data[ci] if ci < len(row_data) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = font_size
            if ri == 0:
                run.bold = True

    doc.add_paragraph()  # spacer after table


def _generate_appendix(analysis, output_path, font_name):
    """Generate appendix .docx matching the reference appendix format."""
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from generate_docx_report import generate_profile_chart

    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    # Sort entries
    all_entries = []
    for key, result in analysis.get("flights", {}).items():
        parts = key.rsplit("_", 1)
        all_entries.append((parts[0], parts[1] if len(parts) > 1 else "", result))

    for flight, date_str, result in sorted(all_entries):
        # Simple header: "4月2日 I99806"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{date_str}  {flight}")
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(12)
        run.bold = True

        # Generate chart
        chart_dir = os.path.join(REPORTS_DIR, "charts")
        os.makedirs(chart_dir, exist_ok=True)
        chart_path = os.path.join(chart_dir, f"{key}_profile.png")

        plan_file, actual_file = find_csv_files(flight, date_str)
        if plan_file and actual_file:
            plan_wp = load_plan_track(plan_file)
            actual_pts = load_actual_track(actual_file)
            ci = CountryIndex(GEOJSON_PATH)
            dev_results = compute_deviations(actual_pts, plan_wp,
                {"min_alt_deviation_ft": 300, "min_duration_nm": 30}, ci)
            mark_cruise_points(dev_results, plan_wp)
            generate_profile_chart(plan_wp, actual_pts, dev_results, flight, date_str, chart_path)
            if os.path.exists(chart_path):
                doc.add_picture(chart_path, width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # No text notes in appendix per user request

    doc.save(output_path)


# ─── Main CLI ─────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="AI航班偏差分析报告生成器")
    parser.add_argument("--flight", "-f", help="航班号(逗号分隔)")
    parser.add_argument("--dates", help="日期列表，逗号分隔")
    parser.add_argument("--route", default=None, help="航线名称")
    parser.add_argument("--min-alt-dev", type=float, default=DEFAULT_MIN_ALT_DEV_FT)
    parser.add_argument("--min-dur", type=float, default=DEFAULT_MIN_DURATION_NM)
    parser.add_argument("--data", help="已有 analysis.json")
    parser.add_argument("--json", help="导出 analysis.json 路径")
    parser.add_argument("--reference", help="参考主报告 .docx")
    parser.add_argument("--appendix-ref", help="参考附件 .docx")
    parser.add_argument("--prompt-only", action="store_true")
    parser.add_argument("--prompt", help="自定义 prompt 文件")
    parser.add_argument("--output", "-o", default=None)
    parser.add_argument("--finalize", action="store_true")
    parser.add_argument("--draft", help="已审核 draft markdown 文件")

    args = parser.parse_args()
    api_config = load_api_config()

    # ── Step 1: Get analysis data ──
    analysis = None
    if args.data:
        with open(args.data, "r", encoding="utf-8") as f:
            analysis = json.load(f)
        print(f"加载分析数据: {args.data}")
    elif args.flight and args.dates:
        flights = [f.strip() for f in args.flight.split(",")]
        dates = [d.strip() for d in args.dates.split(",")]
        route = args.route or f"{flights[0]}航线"

        print("加载国界数据...")
        country_index = CountryIndex(GEOJSON_PATH)
        config = {
            "min_alt_deviation_ft": args.min_alt_dev,
            "min_duration_nm": args.min_dur,
            "route_name": route,
            "date_range": f"{dates[0]}至{dates[-1]}",
        }

        print(f"\n分析航班: {', '.join(flights)}, 日期: {', '.join(dates)}")
        flights_data = {}
        for flight_num in flights:
            for date_str in dates:
                key = f"{flight_num}_{date_str}"
                print(f"  分析: {key}...")
                pf, af = find_csv_files(flight_num, date_str)
                if not pf or not af:
                    print(f"    SKIP: 找不到文件")
                    continue
                result = analyze_flight(pf, af, config, country_index)
                if result:
                    flights_data[key] = result

        analysis = {"flights": flights_data, "analysis_time": datetime.now().isoformat()}
        print(f"\n共分析 {len(flights_data)} 个航班记录")

        if args.json:
            with open(args.json, "w", encoding="utf-8") as f:
                json.dump(analysis, f, ensure_ascii=False, indent=2)
            print(f"分析数据已导出: {args.json}")
    else:
        print("ERROR: 请指定 --data 或 --flight + --dates"); sys.exit(1)

    # ── Paths ──
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    route_name = (args.route or "report").replace("/", "-")
    draft_path = args.output or os.path.join(REPORTS_DIR, f"{route_name}_draft_{timestamp}.md")
    prompt_path = draft_path.replace(".md", "_prompt.txt")

    # ── Finalize mode ──
    if args.finalize:
        draft_file = args.draft or draft_path
        if not os.path.exists(draft_file):
            print(f"ERROR: {draft_file} not found"); sys.exit(1)
        docx_path = draft_file.replace(".md", ".docx")
        finalize_to_docx(draft_file, analysis, docx_path)
        return

    # ── Build prompts ──
    def _read_docx(path):
        if not path or not os.path.exists(path): return ""
        if path.endswith(".docx"):
            from docx import Document
            doc = Document(path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    txt = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                    if txt: parts.append(txt)
            return "\n\n".join(parts)
        with open(path, encoding="utf-8") as f: return f.read()

    ref_text = _read_docx(args.reference)
    appendix_text = _read_docx(args.appendix_ref)

    system_prompt = build_system_prompt(ref_text, appendix_text)
    user_prompt = build_user_prompt(analysis, {
        "route_name": args.route or "未指定",
        "date_range": getattr(args, 'date_range', ''),
    })

    full_prompt = f"=== SYSTEM PROMPT ===\n{system_prompt}\n\n=== USER PROMPT ===\n{user_prompt}"
    with open(prompt_path, "w", encoding="utf-8") as f:
        f.write(full_prompt)
    print(f"\nPrompt 已保存: {prompt_path}")

    if args.prompt_only:
        print("\n" + "=" * 60)
        print(full_prompt)
        print("=" * 60)
        return

    # ── Load custom prompt ──
    if args.prompt:
        with open(args.prompt, encoding="utf-8") as f:
            custom = f.read()
        if "=== USER PROMPT ===" in custom:
            parts = custom.split("=== USER PROMPT ===")
            system_prompt = parts[0].replace("=== SYSTEM PROMPT ===\n", "").strip()
            user_prompt = parts[1].strip()
        else:
            user_prompt = custom

    # ── Call API ──
    print(f"\n调用 DeepSeek API (model: {api_config['model']})...")
    try:
        response = call_deepseek([
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ], api_config)
    except Exception as e:
        print(f"API 调用失败: {e}")
        print(f"Prompt 已保存至 {prompt_path}")
        sys.exit(1)

    with open(draft_path, "w", encoding="utf-8") as f:
        f.write(response)
    print(f"AI 初稿已生成: {draft_path}")
    print(f"\n审阅后: python review_report.py --data ... --draft {draft_path} --finalize")


if __name__ == "__main__":
    main()
